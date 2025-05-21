
import sys
import os
import time
import threading
import sqlite3
import hashlib
from concurrent.futures import ThreadPoolExecutor

from PyQt5.QtCore import (QCoreApplication, QThread, pyqtSignal, Qt, QWaitCondition, QMutex, QPropertyAnimation, QRect, QTimer, QEasingCurve, QPoint)
from PyQt5.QtGui import QIcon, QFont, QPixmap, QColor, QPainter, QLinearGradient, QPalette, QBrush, QRegion, QPolygon, QPainterPath, QMovie
from PyQt5.QtWidgets import (QStyle, QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGridLayout, QLabel, QLineEdit, QPushButton, QFileDialog, QListWidget, QComboBox, QMessageBox, QProgressBar, QDialog, QTableWidget, QTableWidgetItem, QHeaderView, QInputDialog, QGraphicsDropShadowEffect, QGroupBox, QListView, QTreeView, QFrame, QStackedWidget, QGraphicsOpacityEffect)

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph

from docx import Document
from docx.shared import Pt

import matplotlib
matplotlib.use("Qt5Agg")
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure

# ---------------- SplashScreen for Animated Start-up ----------------
class SplashScreen(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent, Qt.FramelessWindowHint)
        self.setModal(True)
        self.setFixedSize(600, 300)
        # Set a light purple gradient background for the splash screen
        palette = QPalette()
        gradient = QLinearGradient(0, 0, self.width(), 0)
        gradient.setColorAt(0, QColor("#E6E6FA"))  # Lavender
        gradient.setColorAt(0.5, QColor("#D8BFD8"))  # Thistle
        gradient.setColorAt(1, QColor("#DDA0DD"))    # Plum
        palette.setBrush(QPalette.Window, QBrush(gradient))
        self.setPalette(palette)
        self.setAutoFillBackground(True)

        layout = QVBoxLayout(self)
        self.label_animation = QLabel("", self)
        self.label_animation.setAlignment(Qt.AlignCenter)
        font = QFont("Segoe Script", 24, QFont.Bold)

        self.label_animation.setFont(font)
        self.label_animation.setStyleSheet("color: #4B0082;")  # Indigo text
        layout.addWidget(self.label_animation)
        self.full_text = "التحقيق الجنائي الرقمي\nDigital Forensic Investigation"
        self.current_index = 0

        # Timer for typewriter effect
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_text)
        self.timer.start(100)

    def update_text(self):
        if self.current_index < len(self.full_text):
            self.label_animation.setText(self.label_animation.text() + self.full_text[self.current_index])
            self.current_index += 1
        else:
            self.timer.stop()
            # After complete, wait briefly and then fade out
            QTimer.singleShot(1000, self.start_fade_out)

    def start_fade_out(self):
        self.effect = QGraphicsOpacityEffect(self)
        self.setGraphicsEffect(self.effect)
        self.fade_anim = QPropertyAnimation(self.effect, b"opacity")
        self.fade_anim.setDuration(1000)
        self.fade_anim.setStartValue(1)
        self.fade_anim.setEndValue(0)
        self.fade_anim.setEasingCurve(QEasingCurve.InOutQuad)
        self.fade_anim.finished.connect(self.close)
        self.fade_anim.start()

# ---------------- HoverButton Class for Visual Hover Effects ----------------
class HoverButton(QPushButton):
    def __init__(self, text="", parent=None):
        super().__init__(text, parent)
        self.default_shadow = QGraphicsDropShadowEffect(self)
        self.default_shadow.setBlurRadius(5)
        self.default_shadow.setColor(QColor(0, 0, 0, 160))
        self.default_shadow.setOffset(0, 0)
        self.setGraphicsEffect(self.default_shadow)
        self.setStyleSheet("""
            QPushButton {
                border: none;
                padding: 8px;
            }
            QPushButton:hover {
                background-color: rgba(255, 255, 255, 30);
            }
        """)

    def enterEvent(self, event):

        self.default_shadow.setBlurRadius(15)
        super().enterEvent(event)

    def leaveEvent(self, event):
        self.default_shadow.setBlurRadius(5)
        super().leaveEvent(event)

# ---------------- HourglassWidget for Animated Hourglass ----------------
class HourglassWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.angle = 0
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_angle)
        self.timer.setInterval(100)  # update every 100 ms
        self.timer.start()
        self.setFixedSize(40, 40)

    def update_angle(self):
        self.angle = (self.angle + 30) % 360
        self.update()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        center = self.rect().center()
        painter.translate(center)
        painter.rotate(self.angle)
        painter.translate(-center)
        path = QPainterPath()
        w = self.width()
        h = self.height()
        path.moveTo(w/2, 0)
        path.lineTo(0, h/2)
        path.lineTo(w, h/2)
        path.closeSubpath()
        path.moveTo(0, h/2)
        path.lineTo(w/2, h)
        path.lineTo(w, h/2)
        path.closeSubpath()
        painter.setBrush(QColor("#FFA500"))
        painter.drawPath(path)

# ---------------- Database Manager ----------------
class DatabaseManager:
    def __init__(self):
        self.conn = sqlite3.connect('file_search.db', check_same_thread=False)
        self._init_db()

    def _init_db(self):
        tables = {

            'search_history': '''
                CREATE TABLE IF NOT EXISTS search_history (
                    id INTEGER PRIMARY KEY,
                    file_path TEXT UNIQUE,
                    file_hash TEXT,
                    extension TEXT,
                    search_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''',
            'non_matching_hashes': '''
                CREATE TABLE IF NOT EXISTS non_matching_hashes (
                    id INTEGER PRIMARY KEY,
                    file_path TEXT UNIQUE,
                    file_hash TEXT,
                    extension TEXT,
                    search_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            '''
        }
        with self.conn:
            for schema in tables.values():
                self.conn.execute(schema)

    def save_record(self, table_name, file_path, file_hash, extension):
        try:
            with self.conn:
                self.conn.execute(
                    f'''INSERT OR IGNORE INTO {table_name}
                    (file_path, file_hash, extension)
                    VALUES (?, ?, ?)''',
                    (file_path, file_hash, extension)
                )
        except sqlite3.Error as e:
            raise Exception(f"Database error: {str(e)}")

    def search_hash(self, target_hash):
        try:
            with self.conn:
                cursor = self.conn.execute('''
                    SELECT file_path, file_hash FROM search_history
                    WHERE file_hash = ?
                    UNION ALL
                    SELECT file_path, file_hash FROM non_matching_hashes
                    WHERE file_hash = ?
                ''', (target_hash, target_hash))
                return cursor.fetchall()
        except sqlite3.Error as e:
            raise Exception(f"Database error: {str(e)}")

    def search_non_matching(self, target_hash):
        try:

            with self.conn:
                cursor = self.conn.execute('''
                    SELECT file_path, file_hash FROM non_matching_hashes
                    WHERE file_hash = ?
                ''', (target_hash,))
                return cursor.fetchall()
        except sqlite3.Error as e:
            raise Exception(f"Database error: {str(e)}")

    def delete_record(self, file_path, file_hash):
        try:
            with self.conn:
                self.conn.execute("DELETE FROM search_history WHERE file_path=? AND file_hash=?", (file_path, file_hash))
                self.conn.execute("DELETE FROM non_matching_hashes WHERE file_path=? AND file_hash=?", (file_path, file_hash))
        except sqlite3.Error as e:
            raise Exception(f"Database delete error: {str(e)}")

# ---------------- Base Search Thread ----------------
class BaseSearchThread(QThread):
    progress_updated = pyqtSignal(int)  # percentage; -1 indicates indeterminate progress
    result_found = pyqtSignal(str, str)
    error_occurred = pyqtSignal(str)
    finished = pyqtSignal()

    def __init__(self):
        super().__init__()
        self._is_paused = False
        self._is_stopped = False
        self.condition = QWaitCondition()
        self.mutex = QMutex()

    def pause(self):
        self._is_paused = True

    def resume(self):
        self._is_paused = False
        self.condition.wakeAll()

    def stop(self):
        self._is_stopped = True
        self.resume()

# ---------------- Local Search Thread with Enhanced Non-Blocking Scanning ----------------
class LocalSearchThread(BaseSearchThread):
    def __init__(self, paths, target_hash, extensions, excluded_paths):
        super().__init__()
        self.paths = paths
        self.target_hash = target_hash
        self.extensions = [ext for ext in extensions if ext != "all"]

        self.excluded_paths = [os.path.normpath(p) for p in excluded_paths]
        self.db = DatabaseManager()
        self._pause_event = threading.Event()
        self._pause_event.set()

    def _should_exclude(self, current_path):
        current = os.path.normpath(current_path)
        return any(os.path.commonpath([current, excluded]) == excluded for excluded in self.excluded_paths)

    def scan_directory(self, path):
        try:
            with os.scandir(path) as it:
                for entry in it:
                    if self._is_stopped:
                        break
                    if entry.is_dir(follow_symlinks=False):
                        if self._should_exclude(entry.path):
                            continue
                        yield from self.scan_directory(entry.path)
                    elif entry.is_file(follow_symlinks=False):
                        if self.extensions and not entry.name.endswith(tuple(self.extensions)):
                            continue
                        yield entry.path
        except Exception:
            return

    def process_file(self, file_path):
        if self._is_stopped:
            return
        self._pause_event.wait()
        try:
            with open(file_path, 'rb') as f:
                hasher = hashlib.sha256()
                while True:
                    chunk = f.read(131072)
                    if not chunk:
                        break
                    hasher.update(chunk)
                file_hash = hasher.hexdigest()
        except Exception:
            return
        if file_hash == self.target_hash:
            self.result_found.emit(file_path, file_hash)
            self.mutex.lock()
            try:
                self.db.save_record('search_history', file_path, file_hash, os.path.splitext(file_path)[1])
            finally:
                self.mutex.unlock()
        else:
            self.mutex.lock()

            try:
                self.db.save_record('non_matching_hashes', file_path, file_hash, os.path.splitext(file_path)[1])
            finally:
                self.mutex.unlock()

    def run(self):
        try:
            processed_count = 0
            self.progress_updated.emit(-1)
            max_workers = os.cpu_count() or 4
            futures = []

            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                for base_path in self.paths:
                    for file_path in self.scan_directory(base_path):
                        if self._is_stopped:
                            break
                        futures.append(executor.submit(self.process_file, file_path))
                        processed_count += 1
                        if processed_count % 50 == 0:
                            QCoreApplication.processEvents()
                for future in futures:
                    if self._is_stopped:
                        break
                    future.result()
            self.finished.emit()
        except Exception as e:
            self.error_occurred.emit(f"Critical error: {str(e)}")

# ---------------- History Loader Thread ----------------
class HistoryLoaderThread(QThread):
    history_loaded = pyqtSignal(list)
    def __init__(self, db):
        super().__init__()
        self.db = db
    def run(self):
        try:
            query = '''
                SELECT search_date, extension, file_hash, file_path
                FROM search_history
                ORDER BY search_date DESC
            '''
            with self.db.conn:
                cursor = self.db.conn.execute(query)
                data = cursor.fetchall()
                records = [f"{date}::{ext}::{hash_val}::{path}" for date, ext, hash_val, path in data]
                self.history_loaded.emit(records)
        except Exception as e:
            self.history_loaded.emit([f"Error loading history: {str(e)}"])

# ---------------- Smart Check Thread for Non-Matching Hashes ----------------

class SmartCheckThread(QThread):
    result_ready = pyqtSignal(list)
    def __init__(self, db, target_hash):
        super().__init__()
        self.db = db
        self.target_hash = target_hash
    def run(self):
        try:
            results = self.db.search_non_matching(self.target_hash)
            self.result_ready.emit(results)
        except Exception as e:
            self.result_ready.emit([])

# ---------------- Contemporary Chart Widget ----------------
class ContemporaryChartWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.figure = Figure(figsize=(5, 4))
        self.canvas = FigureCanvas(self.figure)
        layout = QVBoxLayout()
        layout.addWidget(self.canvas)
        self.setLayout(layout)
        self.disk_count = 0
        self.smart_count = 0
        self.draw_chart(0, 0)

    def update_chart(self, disk_count, smart_count):
        self.disk_count = disk_count
        self.smart_count = smart_count
        self.draw_chart(self.disk_count, self.smart_count)

    def draw_chart(self, disk, smart):
        self.figure.clear()
        ax = self.figure.add_subplot(111)
        total = disk + smart
        if total == 0:
            data = [1]
            labels = ['No Data']
            colors = ['#D3D3D3']
        else:
            data = [disk, smart]
            labels = ['Disk Search', 'Smart Search']
            colors = ['#4CAF50', '#03A9F4']
        ax.pie(data, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors,
               wedgeprops=dict(width=0.4, edgecolor='w'))
        ax.set_title("Search Results", fontsize=16, fontweight='bold')
        ax.axis('equal')
        self.canvas.draw()

# ---------------- Triangle Widget for Flip Animation (New Feature) ----------------
class TriangleWidget(QWidget):

    def __init__(self, front_text, back_text, front_color=None, back_color=None, parent=None):
        super().__init__(parent)
        self.front_text = front_text
        self.back_text = back_text
        self.front_color = front_color if front_color else "#0072ff"
        self.back_color = back_color if back_color else "#ff416c"
        self.current_state = 0  # 0: front, 1: back
        self.label = QLabel(self.front_text, self)
        self.label.setAlignment(Qt.AlignCenter)
        font = QFont("Arial", 10, QFont.Bold)
        self.label.setFont(font)
        self.opacity_effect = QGraphicsOpacityEffect(self.label)
        self.label.setGraphicsEffect(self.opacity_effect)

        self.anim = QPropertyAnimation(self.opacity_effect, b"opacity")
        self.anim.setEasingCurve(QEasingCurve.InOutQuad)
        self.anim.setDuration(500)
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.flip)
        self.timer.start(3000)  # flip every 3 seconds
        self.setMinimumWidth(120)
        self.setMinimumHeight(100)
        self.setMaximumHeight(120)
        self.setAutoFillBackground(False)

    def resizeEvent(self, event):
        w = self.width()
        h = self.height()
        points = QPolygon([
            QPoint(w // 2, 0),
            QPoint(0, h),
            QPoint(w, h)
        ])
        self.setMask(QRegion(points))
        self.label.setGeometry(int(w * 0.2), int(h * 0.4), int(w * 0.6), int(h * 0.4))
        super().resizeEvent(event)

    def paintEvent(self, event):
        painter = QPainter(self)
        rect = self.rect()
        gradient = QLinearGradient(0, 0, rect.width(), rect.height())
        if self.current_state == 0:
            gradient.setColorAt(0, QColor(self.front_color))
            gradient.setColorAt(1, QColor("#00c6ff"))
        else:
            gradient.setColorAt(0, QColor(self.back_color))
            gradient.setColorAt(1, QColor("#ff4b2b"))
        painter.setBrush(QBrush(gradient))
        painter.setPen(Qt.NoPen)
        points = QPolygon([
            QPoint(rect.width() // 2, 0),

            QPoint(0, rect.height()),
            QPoint(rect.width(), rect.height())
        ])
        painter.drawPolygon(points)
        super().paintEvent(event)

    def flip(self):
        self.anim.stop()
        self.anim.setStartValue(1.0)
        self.anim.setEndValue(0.0)
        self.anim.finished.connect(self.toggle_content)
        self.anim.start()

    def toggle_content(self):
        self.anim.finished.disconnect(self.toggle_content)
        self.current_state = 1 - self.current_state
        if self.current_state == 0:
            self.label.setText(self.front_text)
        else:
            self.label.setText(self.back_text)
        self.anim.setStartValue(0.0)
        self.anim.setEndValue(1.0)
        self.anim.start()

# ---------------- Triangle Chain Widget (New Feature) ----------------
class TriangleChainWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.layout = QHBoxLayout()
        self.layout.setSpacing(10)
        self.setLayout(self.layout)
        self.build_triangles()

    def build_triangles(self):
        items = [
            {"front": "البحث الذكي", "back": "Smart & Efficient", "front_color": "#FF5733", "back_color": "#FFC300"},
            {"front": "توقيع رقمي", "back": "Secure Signature", "front_color": "#33FF57", "back_color": "#33FFF3"},
            {"front": "تصدير سريع", "back": "Fast Export", "front_color": "#3357FF", "back_color": "#8C33FF"},
            {"front": "آمن وموثوق", "back": "Reliable & Safe", "front_color": "#FF33A8", "back_color": "#FF5733"},
            {"front": "تحليل سريع", "back": "Quick Analysis", "front_color": "#33FFA8", "back_color": "#33A8FF"},
            {"front": "تحقيق جنائي", "back": "Forensic Investigation", "front_color": "#A833FF", "back_color": "#FF33F6"}
        ]
        for item in items:
            if item["front"] == "تحقيق جنائي":
                container = QWidget()
                hlayout = QHBoxLayout()
                hlayout.setSpacing(5)
                hlayout.setContentsMargins(0, 0, 0, 0)
                triangle = TriangleWidget(item["front"], item["back"], item["front_color"], item["back_color"])
                hourglass = HourglassWidget()

                hlayout.addWidget(triangle)
                hlayout.addWidget(hourglass)
                container.setLayout(hlayout)
                self.layout.addWidget(container)
            else:
                triangle = TriangleWidget(item["front"], item["back"], item["front_color"], item["back_color"])
                self.layout.addWidget(triangle)
        self.layout.addStretch()

# ---------------- File Database Dialog (Non-Matching Hashes) ----------------
class FileDatabaseDialog(QDialog):
    def __init__(self, db):
        super().__init__()
        self.db = db
        self.setWindowTitle("Non-Matching Hash Records")
        self.setGeometry(200, 200, 1000, 600)
        self.init_ui()
        self.apply_dialog_style()
        threading.Thread(target=self.load_data, daemon=True).start()

    def init_ui(self):
        layout = QVBoxLayout()
        self.search_line_edit = QLineEdit()
        self.search_line_edit.setPlaceholderText("Search records...")
        self.search_line_edit.textChanged.connect(self.filter_table)
        layout.addWidget(self.search_line_edit)
        self.table = QTableWidget()
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["Date", "Extension", "Digital Signature", "Path"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.table)
        self.setLayout(layout)

    def load_data(self):
        try:
            query = '''
                SELECT search_date, extension, file_hash, file_path
                FROM non_matching_hashes
                ORDER BY search_date DESC
            '''
            with self.db.conn:
                cursor = self.db.conn.execute(query)
                data = cursor.fetchall()
            self.table.setRowCount(len(data))
            for row, (date, ext, hash_val, path) in enumerate(data):
                self.table.setItem(row, 0, QTableWidgetItem(date))
                self.table.setItem(row, 1, QTableWidgetItem(ext))
                self.table.setItem(row, 2, QTableWidgetItem(hash_val))
                self.table.setItem(row, 3, QTableWidgetItem(path))
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error loading data: {str(e)}")


    def filter_table(self, text):
        text = text.lower().strip()
        for row in range(self.table.rowCount()):
            match = False
            for col in range(self.table.columnCount()):
                item = self.table.item(row, col)
                if item and text in item.text().lower():
                    match = True
                    break
            self.table.setRowHidden(row, not match)

    def apply_dialog_style(self):
        self.setStyleSheet("""
            QDialog {
                background: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:1,
                              stop:0 #8E2DE2, stop:1 #4A00E0);
            }
            QLineEdit {
                background-color: #ffffff;
                border-radius: 6px;
                padding: 6px;
            }
            QTableWidget {
                background-color: #ffffff;
            }
            QHeaderView::section {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #FF8C00, stop:1 #FFA500);
                color: white;
                font-weight: bold;
            }
        """)

# ---------------- History Dialog ----------------
class HistoryDialog(QDialog):
    def __init__(self, db):
        super().__init__()
        self.db = db
        self.setWindowTitle("History")
        self.setGeometry(300, 300, 800, 500)
        self.init_ui()
        self.apply_dialog_style()
        self.load_history_async()

    def init_ui(self):
        main_layout = QVBoxLayout()
        top_layout = QHBoxLayout()
        self.search_line_edit = QLineEdit()
        self.search_line_edit.setPlaceholderText("Search in table...")
        self.search_line_edit.textChanged.connect(self.filter_table)
        top_layout.addWidget(self.search_line_edit)

        self.filter_combo = QComboBox()
        self.filter_combo.addItems(["All", ".txt", ".pdf", ".docx", ".jpg", ".exe", ".sys", ".sdb", ".ioc", ".dll", ".xml"])
        self.filter_combo.currentTextChanged.connect(self.filter_by_extension)
        top_layout.addWidget(QLabel("Filter by Extension:"))
        top_layout.addWidget(self.filter_combo)
        main_layout.addLayout(top_layout)
        self.table = QTableWidget()
        self.table.setColumnCount(9)
        self.table.setHorizontalHeaderLabels([
            "Name", "Path", "Source", "File Type", "Date", "Digital Signature", "Age", "Frequency", "User"
        ])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        main_layout.addWidget(self.table)
        bottom_layout = QHBoxLayout()
        self.btn_delete = HoverButton("Delete Selected")
        self.btn_export_selected = HoverButton("Export Selected")
        bottom_layout.addWidget(self.btn_delete)
        bottom_layout.addWidget(self.btn_export_selected)
        main_layout.addLayout(bottom_layout)
        self.setLayout(main_layout)
        self.btn_delete.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_TrashIcon', QStyle.SP_DialogCancelButton)))
        self.btn_export_selected.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_DialogSaveButton', QStyle.SP_DialogOkButton)))
        self.btn_delete.clicked.connect(self.delete_selected_rows)
        self.btn_export_selected.clicked.connect(self.export_selected_rows)

    def apply_dialog_style(self):
        self.setStyleSheet("""
            QDialog {
                background: qlineargradient(spread:pad, x1:0, y1:0, x2:1, y2:1,
                              stop:0 #F7971E, stop:1 #FFD200);
            }
            QLineEdit {
                background-color: #ffffff;
                border-radius: 6px;
                padding: 6px;
            }
            QTableWidget {
                background-color: #ffffff;
            }
            QHeaderView::section {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #1E90FF, stop:1 #00BFFF);
                color: white;
                font-weight: bold;
            }
            QPushButton {
                font-weight:bold;
                font-size:14px;
                border-radius: 8px;
                padding: 8px;

            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1, stop:0 #ffd700, stop:1 #ff8c00);
            }
        """)

    def load_history_async(self):
        self.table.setRowCount(0)
        self.loader_thread = HistoryLoaderThread(self.db)
        self.loader_thread.history_loaded.connect(self.populate_history)
        self.loader_thread.start()

    def populate_history(self, records):
        self.table.setRowCount(len(records))
        for row_index, rec in enumerate(records):
            if "::" in rec:
                date, ext, hash_val, path = rec.split("::")
                name = os.path.basename(path) if path else "N/A"
                source = "Local DB"
                file_type = ext
                digital_sig = hash_val
                the_date = date
                try:
                    creation_time = os.path.getctime(path)
                    age_days = (time.time() - creation_time) / 86400.0
                    age = f"{age_days:.1f} days"
                except:
                    age = "N/A"
                frequency = "1"
                user = "System" if "windows" in path.lower() else "User"
                self.table.setItem(row_index, 0, QTableWidgetItem(name))
                self.table.setItem(row_index, 1, QTableWidgetItem(path))
                self.table.setItem(row_index, 2, QTableWidgetItem(source))
                self.table.setItem(row_index, 3, QTableWidgetItem(file_type))
                self.table.setItem(row_index, 4, QTableWidgetItem(the_date))
                self.table.setItem(row_index, 5, QTableWidgetItem(digital_sig))
                self.table.setItem(row_index, 6, QTableWidgetItem(str(age)))
                self.table.setItem(row_index, 7, QTableWidgetItem(frequency))
                self.table.setItem(row_index, 8, QTableWidgetItem(user))
            else:
                self.table.setRowCount(1)
                self.table.setItem(0, 0, QTableWidgetItem(rec))

    def filter_table(self, text):
        text = text.lower().strip()
        for row in range(self.table.rowCount()):
            row_hidden = True
            for col in range(self.table.columnCount()):
                item = self.table.item(row, col)
                if item and text in item.text().lower():
                    row_hidden = False

                    break
            self.table.setRowHidden(row, row_hidden)

    def filter_by_extension(self, ext):
        ext = ext.lower().strip()
        for row in range(self.table.rowCount()):
            item_ext = self.table.item(row, 3)
            if item_ext:
                if ext == "all":
                    self.table.setRowHidden(row, False)
                else:
                    self.table.setRowHidden(row, not (ext == item_ext.text().lower().strip()))

    def delete_selected_rows(self):
        selected_ranges = self.table.selectedRanges()
        if not selected_ranges:
            QMessageBox.information(self, "Info", "No rows selected for deletion.")
            return
        rows_to_delete = []
        for sel_range in selected_ranges:
            for row in range(sel_range.topRow(), sel_range.bottomRow() + 1):
                rows_to_delete.append(row)
        rows_to_delete = sorted(set(rows_to_delete), reverse=True)

        for row in rows_to_delete:
            path_item = self.table.item(row, 1)
            hash_item = self.table.item(row, 5)
            if path_item and hash_item:
                path = path_item.text()
                hash_val = hash_item.text()
                try:
                    self.db.delete_record(path, hash_val)
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Error deleting record: {str(e)}")
            self.table.removeRow(row)

    def export_selected_rows(self):
        selected_ranges = self.table.selectedRanges()
        if not selected_ranges:
            QMessageBox.information(self, "Info", "No rows selected for export.")
            return
        rows_to_export = []
        for sel_range in selected_ranges:
            for row in range(sel_range.topRow(), sel_range.bottomRow() + 1):
                rows_to_export.append(row)
        rows_to_export = sorted(set(rows_to_export))
        fmt, ok = QInputDialog.getItem(
            self, "Export Format", "Select format:", ["PDF", "Word", "Excel"], 0, False
        )
        if not ok:
            return

        if fmt == "Excel":
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Save File", "", "Excel Files (*.xlsx)"
            )
        else:
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Save File", "", f"{fmt} Files (*.{fmt.lower()})"
            )
        if not file_path:
            return
        exported_data = []
        headers = ["Name", "Path", "Source", "File Type", "Date", "Digital Signature", "Age", "Frequency", "User"]
        exported_data.append(headers)
        for row in rows_to_export:
            row_data = []
            for col in range(self.table.columnCount()):
                item = self.table.item(row, col)
                row_data.append(item.text() if item else "")
            exported_data.append(row_data)
        try:
            if fmt == "PDF":
                self.export_to_pdf(file_path, exported_data)
            elif fmt == "Word":
                self.export_to_word(file_path, exported_data)
            else:
                self.export_to_excel(file_path, exported_data)
            QMessageBox.information(self, "Success", f"Exported successfully to:\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Export failed: {str(e)}")

    def export_to_pdf(self, filename, data):
        c = canvas.Canvas(filename, pagesize=letter)
        width, height = letter
        style = ParagraphStyle(
            name='Normal',
            fontName='Arial',
            fontSize=10,
            alignment=0,
            spaceAfter=6
        )
        c.setFont("Arial", 14)
        c.drawString(50, 750, "History - Exported Rows")
        c.line(50, 745, 300, 745)
        y = 720
        for row in data:
            row_str = " | ".join(row)
            p = Paragraph(row_str, style)
            w, h = p.wrap(width - 100, 50)
            if y - h < 50:
                c.showPage()
                y = 750

            p.drawOn(c, 50, y - h)
            y -= h + 6
        c.save()

    def export_to_word(self, filename, data):
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        title = doc.add_paragraph("History - Exported Rows")
        title.alignment = 0
        for row in data:
            row_str = " | ".join(row)
            doc.add_paragraph(row_str)
        doc.save(filename)

    def export_to_excel(self, filename, data):
        try:
            from openpyxl import Workbook
        except ImportError:
            QMessageBox.critical(self, "Error", "openpyxl package is required. Install using 'pip install openpyxl'")
            return
        wb = Workbook()
        ws = wb.active
        for row in data:
            ws.append(row)
        wb.save(filename)

# ---------------- Main Window ----------------
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.db = DatabaseManager()
        self.current_thread = None
        self.excluded_paths = []
        self.dark_mode = False
        self.disk_count = 0
        self.smart_count = 0
        self.setup_stylesheets()
        self.init_ui()
        self.setup_connections()

    def setup_stylesheets(self):
        # Updated light stylesheet with a light purple gradient background
        self.light_stylesheet = """
            QMainWindow {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                     stop:0 #E6E6FA, stop:0.5 #D8BFD8, stop:1 #DDA0DD);
            }
            QGroupBox {

                font: bold 16px;
                border: 2px solid #555;
                border-radius: 8px;
                margin-top: 10px;
                background-color: #ffffff;
                color: #333;
                padding: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 4px;
            }
            QLineEdit, QComboBox, QListWidget {
                background-color: #ffffff;
                color: #333;
                padding: 4px;
                border: 1px solid #ccc;
                border-radius: 4px;
            }
            QPushButton {
                font-weight:bold;
                font-size:14px;
                border-radius: 8px;
                padding: 8px;
            }
            QPushButton:hover {
                background: rgba(255, 255, 255, 40);
            }
        """
        self.dark_stylesheet = """
            QMainWindow { background: #2c2c2c; }
            QGroupBox {
                font: bold 16px;
                border: 2px solid #444;
                border-radius: 8px;
                margin-top: 10px;
                background-color: #3c3c3c;
                color: #fff;
                padding: 10px;
            }
            QLineEdit, QComboBox, QListWidget {
                background-color: #555;
                color: #fff;
                padding: 4px;
                border: 1px solid #666;
                border-radius: 4px;
            }
            QPushButton {
                font-weight:bold;
                font-size:14px;

                border-radius: 8px;
                padding: 8px;
            }
            QPushButton:hover {
                background: rgba(255, 255, 255, 40);
            }
        """

    def init_ui(self):
        self.setWindowTitle("Professional Hash Search")
        self.setGeometry(100, 100, 1200, 800)
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(15)

        # Top bar: Add TriangleChainWidget and exit button
        top_bar = QHBoxLayout()
        self.triangle_chain = TriangleChainWidget()
        top_bar.addWidget(self.triangle_chain)
        top_bar.addStretch()
        self.btn_exit_top = HoverButton("Exit")
        self.btn_exit_top.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_TitleBarCloseButton', QStyle.SP_DialogCancelButton)))
        self.btn_exit_top.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #795548, stop:1 #5D4037); color: white; border-radius: 8px; padding: 8px;")
        top_bar.addWidget(self.btn_exit_top)
        main_layout.addLayout(top_bar)

        # Search folder section
        folder_box = QGroupBox("Select Search Folder 🔍")
        folder_layout = QHBoxLayout()
        self.path_input = QLineEdit()
        self.path_input.setPlaceholderText("Enter search folder path...")
        self.btn_browse = HoverButton("Browse")
        self.btn_browse.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_DirOpenIcon', QStyle.SP_FileIcon)))
        self.btn_browse.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #FF9800, stop:1 #F57C00); color: white; border-radius: 8px; padding: 8px;")
        folder_layout.addWidget(self.path_input)
        folder_layout.addWidget(self.btn_browse)
        folder_box.setLayout(folder_layout)
        main_layout.addWidget(folder_box)

        # Hash calculation section
        hash_box = QGroupBox("SHA-256 Digital Signature 🛡")
        hash_layout = QHBoxLayout()
        self.hash_input = QLineEdit()
        self.hash_input.setPlaceholderText("Enter SHA-256 hash (64 characters)")
        self.btn_calculate = HoverButton("Calculate Hash")
        self.btn_calculate.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_DialogApplyButton',

QStyle.SP_DialogOkButton)))
        self.btn_calculate.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #4CAF50, stop:1 #388E3C); color: white; border-radius: 8px; padding: 8px;")
        hash_layout.addWidget(self.hash_input)
        hash_layout.addWidget(self.btn_calculate)
        hash_box.setLayout(hash_layout)
        main_layout.addWidget(hash_box)

        # Search settings section
        search_box = QGroupBox("Search Settings ⚙️")
        search_layout = QHBoxLayout()
        self.ext_combo = QComboBox()
        self.ext_combo.addItems(["all", ".txt", ".pdf", ".docx", ".jpg", ".exe", ".sys"])
        self.btn_search = HoverButton("Local Search")
        self.btn_search.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_FileDialogContentsView', QStyle.SP_FileIcon)))
        self.btn_search.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2196F3, stop:1 #1976D2); color: white; border-radius: 8px; padding: 8px;")
        self.btn_pause = HoverButton("Pause")
        self.btn_pause.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_MediaPause', QStyle.SP_MediaPlay)))
        self.btn_pause.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #9C27B0, stop:1 #7B1FA2); color: white; border-radius: 8px; padding: 8px;")
        self.btn_resume = HoverButton("Resume")
        self.btn_resume.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_MediaPlay', QStyle.SP_MediaPlay)))
        self.btn_resume.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #00BCD4, stop:1 #0097A7); color: white; border-radius: 8px; padding: 8px;")
        self.btn_cancel = HoverButton("Cancel Search")
        self.btn_cancel.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_DialogCancelButton', QStyle.SP_DialogCancelButton)))
        self.btn_cancel.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #F44336, stop:1 #D32F2F); color: white; border-radius: 8px; padding: 8px;")
        search_layout.addWidget(QLabel("Extension:"))
        search_layout.addWidget(self.ext_combo)
        search_layout.addWidget(self.btn_search)
        search_layout.addWidget(self.btn_pause)
        search_layout.addWidget(self.btn_resume)
        search_layout.addWidget(self.btn_cancel)
        search_box.setLayout(search_layout)
        main_layout.addWidget(search_box)

        # Progress Bar for search
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 0)
        self.progress_bar.hide()
        main_layout.addWidget(self.progress_bar)

        # Process Management Section
        manage_box = QGroupBox("Process Management 🚀")
        manage_layout = QGridLayout()

        self.btn_refresh = HoverButton("Refresh Results")
        self.btn_refresh.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_BrowserReload', QStyle.SP_BrowserReload)))
        self.btn_refresh.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #8BC34A, stop:1 #689F38); color: white; border-radius: 8px; padding: 8px;")
        self.btn_exclude = HoverButton("Exclude Path")
        self.btn_exclude.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_DialogCancelButton', QStyle.SP_DialogCancelButton)))
        self.btn_exclude.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #FFEB3B, stop:1 #FBC02D); color: white; border-radius: 8px; padding: 8px;")
        self.btn_history = HoverButton("History")
        self.btn_history.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_FileDialogInfoView', QStyle.SP_FileIcon)))
        self.btn_history.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3F51B5, stop:1 #303F9F); color: white; border-radius: 8px; padding: 8px;")
        self.btn_file_db = HoverButton("Non-Matching Hash")
        self.btn_file_db.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_DriveFDIcon', QStyle.SP_DriveHDIcon)))
        self.btn_file_db.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #607D8B, stop:1 #455A64); color: white; border-radius: 8px; padding: 8px;")
        self.btn_toggle_theme = HoverButton("Toggle Theme")
        self.btn_toggle_theme.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_BrowserStop', QStyle.SP_BrowserStop)))
        self.btn_toggle_theme.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #E91E63, stop:1 #C2185B); color: white; border-radius: 8px; padding: 8px;")
        manage_layout.addWidget(self.btn_refresh, 0, 0)
        manage_layout.addWidget(self.btn_exclude, 0, 1)
        manage_layout.addWidget(self.btn_history, 1, 0)
        manage_layout.addWidget(self.btn_file_db, 1, 1)
        manage_layout.addWidget(self.btn_toggle_theme, 1, 2)
        manage_box.setLayout(manage_layout)
        main_layout.addWidget(manage_box)

        # Results Display Section with contemporary chart
        results_box = QGroupBox("Search Results 📊")
        results_layout = QHBoxLayout()
        left_results = QVBoxLayout()
        lbl_results = QLabel("Results:")
        self.results_list = QListWidget()
        left_results.addWidget(lbl_results)
        left_results.addWidget(self.results_list)
        self.btn_move_history = HoverButton("Move Selected to History")
        self.btn_move_history.setIcon(self.style().standardIcon(getattr(QStyle, 'SP_ArrowForward', QStyle.SP_ArrowForward)))
        self.btn_move_history.setStyleSheet("background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #03A9F4, stop:1 #0288D1); color: white; border-radius: 8px; padding: 8px;")
        self.btn_move_history.clicked.connect(self.move_results_to_history)
        left_results.addWidget(self.btn_move_history)
        results_layout.addLayout(left_results)
        self.chart = ContemporaryChartWidget()
        results_layout.addWidget(self.chart)

        results_box.setLayout(results_layout)
        main_layout.addWidget(results_box)
        self.setStyleSheet(self.light_stylesheet)

    def setup_connections(self):
        self.btn_browse.clicked.connect(self.browse_folder)
        self.btn_calculate.clicked.connect(self.calculate_hash)
        self.btn_search.clicked.connect(self.start_local_search)
        self.btn_history.clicked.connect(self.show_history)
        self.btn_file_db.clicked.connect(self.show_file_database)
        self.btn_exclude.clicked.connect(self.exclude_path)
        self.btn_pause.clicked.connect(self.pause_search)
        self.btn_resume.clicked.connect(self.resume_search)
        self.btn_cancel.clicked.connect(self.cancel_search)
        self.btn_refresh.clicked.connect(self.refresh_results)
        self.btn_toggle_theme.clicked.connect(self.toggle_theme)
        self.btn_exit_top.clicked.connect(self.exit_program)

    def browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Search Folder")
        if folder:
            self.path_input.setText(folder)

    def calculate_hash(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select a File")
        if file_path:
            try:
                hasher = hashlib.sha256()
                with open(file_path, 'rb') as f:
                    while True:
                        chunk = f.read(131072)
                        if not chunk:
                            break
                        hasher.update(chunk)
                self.hash_input.setText(hasher.hexdigest())
                QMessageBox.information(self, "Success", f"File Hash:\n{hasher.hexdigest()}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Unable to read file: {str(e)}")

    def start_local_search(self):
        target_hash = self.hash_input.text().strip().lower()
        search_path = self.path_input.text()
        if len(target_hash) != 64:
            QMessageBox.warning(self, "Error", "Hash must be 64 characters")
            return
        if not os.path.isdir(search_path):
            QMessageBox.warning(self, "Error", "Invalid search folder")
            return
        self.results_list.clear()
        self.disk_count = 0
        self.smart_count = 0

        self.chart.update_chart(self.disk_count, self.smart_count)
        self.progress_bar.show()
        QMessageBox.information(self, "Info", "Starting search in DB Non-Matching Hash...")
        self.smart_thread = SmartCheckThread(self.db, target_hash)
        self.smart_thread.result_ready.connect(lambda results: self.handle_smart_check_results(results, target_hash, search_path))
        self.smart_thread.start()

    def handle_smart_check_results(self, results, target_hash, search_path):
        if results:
            self.smart_count = len(results)
            for path, hash_val in results:
                if os.path.exists(path):
                    item_text = f"{path} - {hash_val}   - Source: Smart Search 🫠🌸🫠🌸"
                else:
                    item_text = f"[Deleted] {path} - {hash_val}  - Source: Smart Search 🫠🌸"
                self.results_list.addItem(item_text)
            self.chart.update_chart(self.disk_count, self.smart_count)
            self.progress_bar.hide()
            QMessageBox.information(self, "Results", "Found results in DB (Smart Search).")
        else:
            QMessageBox.information(self, "Info", "No results in DB. Starting disk search...")
            self.start_disk_search(target_hash, search_path)

    def start_disk_search(self, target_hash, search_path):
        self.disk_count = 0
        self.current_thread = LocalSearchThread(
            [search_path],
            target_hash,
            [self.ext_combo.currentText()],
            self.excluded_paths
        )
        self.current_thread.result_found.connect(self.handle_result_found)
        self.current_thread.progress_updated.connect(lambda p: None)
        self.current_thread.error_occurred.connect(lambda e: QMessageBox.critical(self, "Error", e))
        self.current_thread.finished.connect(lambda: (
            self.progress_bar.hide(),
            QMessageBox.information(self, "Completed", "Disk search completed")
        ))
        self.current_thread.start()

    def handle_result_found(self, path, hash_val):
        self.disk_count += 1
        self.results_list.addItem(f"{path} - {hash_val} - Source: Disk")
        self.chart.update_chart(self.disk_count, self.smart_count)

    def move_results_to_history(self):
        count = self.results_list.count()
        if count == 0:
            QMessageBox.information(self, "Info", "No results to move.")
            return

        added = 0
        for index in range(count):
            item_text = self.results_list.item(index).text()
            parts = item_text.split(" - ")
            if len(parts) >= 2:
                file_path = parts[0].strip()
                file_hash = parts[1].strip()
                ext = os.path.splitext(file_path)[1]
                try:
                    self.db.save_record('search_history', file_path, file_hash, ext)
                    added += 1
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Error moving record: {str(e)}")
        QMessageBox.information(self, "Done", f"Moved {added} records to History.")

    def show_history(self):
        history_dialog = HistoryDialog(self.db)
        history_dialog.exec_()

    def show_file_database(self):
        file_db_dialog = FileDatabaseDialog(self.db)
        file_db_dialog.exec_()

    def exclude_path(self):
        dialog = QFileDialog(self, "Select Paths to Exclude")
        dialog.setFileMode(QFileDialog.Directory)
        dialog.setOption(QFileDialog.ShowDirsOnly, True)
        dialog.setOption(QFileDialog.DontUseNativeDialog, True)
        for view in dialog.findChildren(QListView):
            view.setSelectionMode(QListView.MultiSelection)
        for tree in dialog.findChildren(QTreeView):
            tree.setSelectionMode(QTreeView.MultiSelection)
        if self.path_input.text():
            dialog.setDirectory(self.path_input.text())
        else:
            dialog.setDirectory(os.path.expanduser("~"))
        if dialog.exec_():
            paths = dialog.selectedFiles()
            if paths:
                for path in paths:
                    norm_path = os.path.normpath(path)
                    if norm_path not in self.excluded_paths:
                        self.excluded_paths.append(norm_path)
                QMessageBox.information(self, "Excluded", f"Paths excluded:\n" + "\n".join(paths))

    def pause_search(self):
        if self.current_thread:
            self.current_thread.pause()
            self.progress_bar.setFormat("Paused - %p%")

    def resume_search(self):

        if self.current_thread:
            self.current_thread.resume()
            self.progress_bar.setFormat("%p%")

    def cancel_search(self):
        if self.current_thread:
            self.current_thread.stop()
            self.progress_bar.hide()
            QMessageBox.information(self, "Cancelled", "Search has been cancelled.")

    def refresh_results(self):
        self.results_list.clear()
        self.chart.draw_chart(0, 0)
        self.disk_count = 0
        self.smart_count = 0
        QMessageBox.information(self, "Refreshed", "Results have been cleared.")

    def toggle_theme(self):
        self.dark_mode = not self.dark_mode
        if self.dark_mode:
            self.setStyleSheet(self.dark_stylesheet)
        else:
            self.setStyleSheet(self.light_stylesheet)
        QMessageBox.information(self, "Theme Toggled", "Theme toggled successfully.")

    def exit_program(self):
        if self.current_thread and self.current_thread.isRunning():
            self.current_thread.stop()
            self.current_thread.wait()
        QApplication.quit()

    def closeEvent(self, event):
        if self.current_thread and self.current_thread.isRunning():
            self.current_thread.stop()
            self.current_thread.wait()
        event.accept()

# ---------------- Main Execution ----------------
if __name__ == "__main__":
    app = QApplication(sys.argv)

    # Show the animated splash screen at startup
    splash = SplashScreen()
    splash.show()
    # Process events to animate splash screen
    app.processEvents()
    # Wait until the splash is closed (it fades out automatically)
    splash.exec_()

    window = MainWindow()
    window.show()

    sys.exit(app.exec_())














































