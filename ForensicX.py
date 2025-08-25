#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Advanced Digital Forensic Investigation Tool -- Enhanced and Refined Version

تم تحسين الواجهة والوظائف مع إضافة لمسات من الفخامة والاحترافية:
١- حذف زر Resume.
٢- حذف تبويب Dashboard.
٣- حذف تبويب Logs وإضافة سجل الأحداث داخل لوحة الإحصائيات بخط صغير.
٤- في قائمة نتائج البحث، عدم مسح النتائج عند بدء بحث جديد بل تضاف النتائج الجديدة إلى النتائج السابقة.
٥- تغيير تسمية "Data Filter (integer)" إلى "عمر الملف" (مثال: آخر تعديل عليه أو تاريخ الإنشاء).
"""

import sys
import os
import time
import threading
import sqlite3
import hashlib
import json
import csv
import datetime

from concurrent.futures import ThreadPoolExecutor

# PyQt5 Imports
from PyQt5.QtCore import (QCoreApplication, QThread, pyqtSignal, Qt, QWaitCondition, QMutex,
                          QPropertyAnimation, QRect, QTimer, QEasingCurve, QSize, QEvent, QUrl)
from PyQt5.QtGui import QIcon, QFont, QPixmap, QColor, QIntValidator
from PyQt5.QtWidgets import (QStyle, QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QGridLayout, QLabel,
                             QLineEdit, QPushButton, QFileDialog, QListWidget, QListWidgetItem, QComboBox, QMessageBox, QProgressBar,
                             QDialog, QTableWidget, QTableWidgetItem, QHeaderView, QInputDialog,
                             QGraphicsDropShadowEffect, QGroupBox, QTabWidget, QTextEdit, QSplitter, QScrollBar)
# For optional media sound effect in splash (if desired)
from PyQt5.QtMultimedia import QSoundEffect

# Matplotlib for statistics charts and interactive dashboard (dashboard tab will be removed)
import matplotlib
matplotlib.use("Qt5Agg")
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure

# For DOCX, PDF export and XLSX export
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt
# XLSX support (requires openpyxl)
try:
    import openpyxl
except ImportError:
    openpyxl = None

# ---------------- Helper: Internal Icon Provider ----------------
def get_icon(name, widget):
    mapping = {
        "search": QStyle.SP_FileDialogContentsView,
        "hash": QStyle.SP_FileIcon,
        "folder": QStyle.SP_DirIcon,
        "settings": QStyle.SP_DesktopIcon,
        "exit": QStyle.SP_DialogCloseButton,
        "stop": QStyle.SP_MediaStop,
        "play": QStyle.SP_MediaPlay,
        "clear": QStyle.SP_DialogResetButton,
        "exclude": QStyle.SP_DirClosedIcon,
        "database": QStyle.SP_DriveDVDIcon,
        "report": QStyle.SP_FileIcon,
        "save": QStyle.SP_DialogSaveButton,
        "refresh": QStyle.SP_BrowserReload,
        "file": QStyle.SP_FileIcon
    }
    std_icon = mapping.get(name, QStyle.SP_FileIcon)
    return widget.style().standardIcon(std_icon)

# ---------------- Custom Animated Progress Bar with Dynamic Status ----------------
class AnimatedProgressBar(QProgressBar):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setRange(0, 100)
        self.setTextVisible(True)
        self.animation = QPropertyAnimation(self, b"value")
        self.animation.setDuration(1000)
        self.animation.setLoopCount(-1)
        self.gradient_offset = 0

        self.timer = QTimer(self)
        self.timer.timeout.connect(self._update_gradient_offset)
        self.timer.start(50)
        self.set_gradient_colors("#6A5ACD", "#8A2BE2")
        self.setStyleSheet("""
            AnimatedProgressBar {
                border: 1px solid #6A5ACD;
                border-radius: 5px;
                text-align: center;
                color: white;
                background-color: rgba(224, 224, 224, 0.7);
            }
            AnimatedProgressBar::chunk {
                border-radius: 5px;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                                            stop:0 #6A5ACD, stop:1 #8A2BE2);
            }
        """)

    def set_gradient_colors(self, color1, color2):
        self._color1 = color1
        self._color2 = color2
        self._update_stylesheet()

    def _update_gradient_offset(self):
        self.gradient_offset = (self.gradient_offset + 0.01) % 1.0
        self._update_stylesheet()

    def _update_stylesheet(self):
        self.setStyleSheet(f"""
            AnimatedProgressBar {{
                border: 1px solid {self._color1};
                border-radius: 5px;
                text-align: center;
                color: white;
                background-color: rgba(224, 224, 224, 0.7);
            }}
            AnimatedProgressBar::chunk {{
                border-radius: 5px;
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                                            stop:0 {self._color1}, stop:1 {self._color2});
            }}
        """)

# ---------------- Enhanced SplashScreen with Advanced Animations and optional sound ----------------
class SplashScreen(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent, Qt.FramelessWindowHint)
        self.setModal(True)
        self.setFixedSize(600, 300)
        self.setAttribute(Qt.WA_TranslucentBackground)

        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setAlignment(Qt.AlignCenter)
        self.bg_widget = QWidget(self)
        # Glassmorphism effect: semi-transparent background with rounded corners
        self.bg_widget.setStyleSheet("""
            QWidget {
                background: rgba(255, 255, 255, 0.2);
                border-radius: 20px;
                border: 1px solid rgba(255, 255, 255, 0.3);
            }
        """)
        self.bg_widget.setLayout(main_layout)

        self.icon_label = QLabel(self.bg_widget)
        icon = get_icon("search", self.icon_label)
        pix = icon.pixmap(96, 96)
        self.icon_label.setPixmap(pix)
        self.icon_label.setAlignment(Qt.AlignCenter)
        main_layout.addWidget(self.icon_label)

        self.title_label = QLabel("", self.bg_widget)
        self.title_label.setAlignment(Qt.AlignCenter)
        self.title_label.setFont(QFont("Segoe UI", 28, QFont.Bold))
        self.title_label.setStyleSheet("color: white;")
        main_layout.addWidget(self.title_label)

        self.full_text = "التحقيق الجنائي الرقمي\nDigital Forensic Investigation"
        self.current_index = 0

        self.dot_label = QLabel("", self.bg_widget)
        self.dot_label.setAlignment(Qt.AlignCenter)
        self.dot_label.setFont(QFont("Arial", 16))
        self.dot_label.setStyleSheet("color: #E0E0E0;")
        main_layout.addWidget(self.dot_label)

        # Optional sound effect on splash screen
        self.sound = QSoundEffect()
        splash_sound_path = os.path.join(os.path.dirname(__file__), "splash.wav")
        if os.path.exists(splash_sound_path):
            self.sound.setSource(QUrl.fromLocalFile(splash_sound_path))
            self.sound.setVolume(0.25)  # 25% volume
            # Uncomment next line if you want to play the sound
            # self.sound.play()

        self.icon_animation = QPropertyAnimation(self.icon_label, b"geometry")
        self.icon_animation.setDuration(1500)
        self.icon_animation.setEasingCurve(QEasingCurve.OutBack)
        start_rect = QRect(self.icon_label.x(), self.icon_label.y() + 50,
                           self.icon_label.width(), self.icon_label.height())
        self.icon_animation.setStartValue(start_rect)
        self.icon_animation.setEndValue(self.icon_label.geometry())
        self.icon_animation.start()

        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_text)
        self.timer.start(70)

        self.dot_timer = QTimer(self)
        self.dot_timer.timeout.connect(self.update_dots)
        self.dot_timer.start(400)
        self.dot_count = 0

        self.fade_out_animation = QPropertyAnimation(self, b"windowOpacity")
        self.fade_out_animation.setDuration(800)
        self.fade_out_animation.setStartValue(1.0)
        self.fade_out_animation.setEndValue(0.0)
        self.fade_out_animation.finished.connect(self.close)

    def update_text(self):
        if self.current_index < len(self.full_text):
            self.title_label.setText(self.full_text[:self.current_index + 1])
            self.current_index += 1
        else:
            self.timer.stop()
            QTimer.singleShot(1000, self.start_fade_out)

    def update_dots(self):
        self.dot_count = (self.dot_count + 1) % 4
        self.dot_label.setText("جار التحميل" + "." * self.dot_count)

    def start_fade_out(self):
        self.dot_timer.stop()
        self.fade_out_animation.start()

# ---------------- HoverButton with Micro-Animations and Glow Effect ----------------
class HoverButton(QPushButton):
    def __init__(self, text="", icon_name=None, parent=None):
        super().__init__(text, parent)
        if icon_name:
            self.setIcon(get_icon(icon_name, self))
            self.setIconSize(QSize(20, 20))
        self.default_shadow = QGraphicsDropShadowEffect(self)
        self.default_shadow.setBlurRadius(8)
        self.default_shadow.setColor(QColor(0, 0, 0, 120))
        self.default_shadow.setOffset(2, 2)
        self.setGraphicsEffect(self.default_shadow)
        self.animation = QPropertyAnimation(self.default_shadow, b"blurRadius")
        self.animation.setDuration(200)
        self.animation.setEasingCurve(QEasingCurve.OutQuad)
        # Initial style with glass-like effect and inner shadow simulation
        self.setStyleSheet("""
            HoverButton {
                border: none;
                padding: 10px 15px;
                border-radius: 8px;
                background-color: rgba(106, 90, 205, 0.8);
                color: white;
                font-weight: bold;
                font-size: 13px;
            }
            HoverButton:hover {
                background-color: rgba(138, 43, 226, 0.9);
            }
            HoverButton:pressed {
                background-color: rgba(91, 60, 155, 0.9);
            }
        """)

    def enterEvent(self, event):
        self.animation.setStartValue(self.default_shadow.blurRadius())
        self.animation.setEndValue(15)
        self.animation.start()
        super().enterEvent(event)

    def leaveEvent(self, event):
        self.animation.setStartValue(self.default_shadow.blurRadius())
        self.animation.setEndValue(8)
        self.animation.start()
        super().leaveEvent(event)

# ---------------- Chart Widget for Statistics with Additional Dashboards ----------------
class ContemporaryChartWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.figure = Figure(figsize=(5, 4), facecolor='none')
        self.canvas = FigureCanvas(self.figure)

        layout = QVBoxLayout()

        # الرسم البياني أولًا
        layout.addWidget(self.canvas)

        # تقليل ارتفاع الـ Logs
        self.log_area = QTextEdit()
        self.log_area.setReadOnly(True)
        self.log_area.setMaximumHeight(40)  # تقليل الارتفاع
        self.log_area.setStyleSheet("font-size: 8pt;")

        layout.addWidget(self.log_area)

        self.setLayout(layout)

        self.disk_count = 0
        self.smart_count = 0
        self.time_line = []  # لتخزين الخط الزمني للرسم البياني
        self.file_counts = []  # لتخزين عدد الملفات المكتشفة مع مرور الوقت

        self.draw_chart(0, 0)

        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(15)
        shadow.setColor(QColor(0, 0, 0, 100))
        shadow.setOffset(4, 4)
        self.setGraphicsEffect(shadow)

    def update_chart(self, disk_count, smart_count):
        self.disk_count = disk_count
        self.smart_count = smart_count
        self.time_line.append(datetime.datetime.now())
        self.file_counts.append(disk_count + smart_count)
        self.draw_chart(self.disk_count, self.smart_count)

    def draw_chart(self, disk, smart):
        self.figure.clear()
        ax1 = self.figure.add_subplot(121)
        total = disk + smart
        if total == 0:
            data = [1]
            labels = ['No Data']
            colors = ['#D3D3D3']
            explode = [0]
        else:
            data = [disk, smart]
            labels = ['Normal Search', 'Smart Search']
            colors = ['#4CAF50', '#03A9F4']
            explode = [0.05 if val == max(data) else 0 for val in data]

        wedges, texts, autotexts = ax1.pie(data, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors,
                                          wedgeprops=dict(width=0.4, edgecolor='w'), explode=explode,
                                          pctdistance=0.85)
        for text in texts:
            text.set_color('black')
            text.set_fontsize(10)
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontsize(10)
            autotext.set_fontweight('bold')
        ax1.set_title("Search Statistics", fontsize=12, fontweight='bold', color='#333333')
        ax1.axis('equal')

        # Line Chart showing files found over time
        ax2 = self.figure.add_subplot(122)
        if self.time_line:
            ax2.plot(self.time_line, self.file_counts, marker='o', linestyle='-', color='#FF9800')
            ax2.set_title("Files Found Over Time", fontsize=12, fontweight='bold', color='#333333')
            ax2.tick_params(axis='x', rotation=45, labelsize=8)
            ax2.set_ylabel("Total Files")
        self.canvas.draw()

    # Method to update log area
    def update_log(self, log_message):
        current_text = self.log_area.toPlainText()
        new_text = current_text + log_message + "\n"
        self.log_area.setPlainText(new_text)
# ---------------- Database Manager ----------------
class DatabaseManager:
    def __init__(self):
        self.conn = sqlite3.connect('file_search.db', check_same_thread=False)
        self._init_db()
    def _init_db(self):
        tables = {
            'search_history': '''
                CREATE TABLE IF NOT EXISTS search_history (
                    id INTEGER PRIMARY KEY,
                    file_path TEXT UNIQUE,
                    file_hash TEXT,
                    extension TEXT,
                    search_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            ''',
            'non_matching_hashes': '''
                CREATE TABLE IF NOT EXISTS non_matching_hashes (
                    id INTEGER PRIMARY KEY,
                    file_path TEXT UNIQUE,
                    file_hash TEXT,
                    extension TEXT,
                    search_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            '''
        }
        with self.conn:
            for schema in tables.values():
                self.conn.execute(schema)
    def save_record(self, table_name, file_path, file_hash, extension):
        try:
            with self.conn:
                self.conn.execute(
                    f'''INSERT OR IGNORE INTO {table_name}
                    (file_path, file_hash, extension)
                    VALUES (?, ?, ?)''',
                    (file_path, file_hash, extension)
                )
        except sqlite3.Error as e:
            raise Exception(f"Database error: {str(e)}")
    def search_hash(self, target_hash):
        try:
            with self.conn:
                cursor = self.conn.execute('''
                    SELECT file_path, file_hash FROM search_history
                    WHERE file_hash = ?
                    UNION ALL
                    SELECT file_path, file_hash FROM non_matching_hashes
                    WHERE file_hash = ?
                ''', (target_hash, target_hash))
                return cursor.fetchall()
        except sqlite3.Error as e:
            raise Exception(f"Database error: {str(e)}")
    def search_non_matching(self, target_hash):
        try:
            with self.conn:
                cursor = self.conn.execute('''
                    SELECT file_path, file_hash FROM non_matching_hashes
                    WHERE file_hash = ?
                ''', (target_hash,))
                return cursor.fetchall()
        except sqlite3.Error as e:
            raise Exception(f"Database error: {str(e)}")
    def delete_record(self, file_path, file_hash):
        try:
            with self.conn:
                self.conn.execute("DELETE FROM search_history WHERE file_path=? AND file_hash=?",
                                  (file_path, file_hash))
                self.conn.execute("DELETE FROM non_matching_hashes WHERE file_path=? AND file_hash=?",
                                  (file_path, file_hash))
        except sqlite3.Error as e:
            raise Exception(f"Database delete error: {str(e)}")

# ---------------- Thread Classes for File Search ----------------
class BaseSearchThread(QThread):
    progress_updated = pyqtSignal(int)
    result_found = pyqtSignal(str, str)
    error_occurred = pyqtSignal(str)
    finished = pyqtSignal()
    def __init__(self):
        super().__init__()
        self._is_paused = False
        self._is_stopped = False
        self.condition = QWaitCondition()
        self.mutex = QMutex()
    def pause(self):
        self._is_paused = True
    def resume(self):
        self._is_paused = False
        self.condition.wakeAll()
    def stop(self):
        self._is_stopped = True
        self.resume()

class LocalSearchThread(BaseSearchThread):
    def __init__(self, paths, target_hash, extensions, excluded_paths, min_size=0, data_filter=None,
                 digital_signature=None):
        super().__init__()
        self.paths = paths
        self.target_hash = target_hash
        self.extensions = [ext for ext in extensions if ext != "all"]
        self.excluded_paths = [os.path.normpath(p) for p in excluded_paths]
        self.db = DatabaseManager()
        self._pause_event = threading.Event()
        self._pause_event.set()
        self.min_size = min_size
        self.data_filter = data_filter
        self.digital_signature = digital_signature
    def _should_exclude(self, current_path):
        current = os.path.normpath(current_path)
        return any(os.path.commonpath([current, excluded]) == excluded for excluded in self.excluded_paths)
    def scan_directory(self, folder):
        try:
            for root, dirs, files in os.walk(folder):
                if self._is_stopped:
                    break
                dirs[:] = [d for d in dirs if not self._should_exclude(os.path.join(root, d))]
                for file in files:
                    if self._is_stopped:
                        break
                    file_path = os.path.join(root, file)
                    if self.extensions and not any(file.lower().endswith(ext.lower()) for ext in self.extensions):
                        continue
                    try:
                        if self.min_size > 0 and os.path.getsize(file_path) < self.min_size:
                            continue
                    except Exception:
                        continue
                    yield file_path
        except Exception:
            return
    def process_file(self, file_path):
        if self._is_stopped:
            return
        self._pause_event.wait()
        try:
            with open(file_path, 'rb') as f:
                hasher = hashlib.sha256()
                while True:
                    chunk = f.read(131072)
                    if not chunk:
                        break
                    hasher.update(chunk)
                file_hash = hasher.hexdigest()
        except Exception:
            return
        if self.digital_signature and self.digital_signature != "All":
            if self.digital_signature not in os.path.basename(file_path):
                pass
        if file_hash == self.target_hash:
            self.result_found.emit(file_path, file_hash)
            self.mutex.lock()
            try:
                self.db.save_record('search_history', file_path, file_hash, os.path.splitext(file_path)[1])
            finally:
                self.mutex.unlock()
        else:
            self.mutex.lock()
            try:
                self.db.save_record('non_matching_hashes', file_path, file_hash, os.path.splitext(file_path)[1])
            finally:
                self.mutex.unlock()
    def run(self):
        try:
            processed_count = 0
            self.progress_updated.emit(-1)
            max_workers = os.cpu_count() or 4
            futures = []
            # تحسين السرعة: تحديث الأحداث كل 10 عملية بدلاً من 25
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                for base_path in self.paths:
                    for file_path in self.scan_directory(base_path):
                        if self._is_stopped:
                            break
                        futures.append(executor.submit(self.process_file, file_path))
                        processed_count += 1
                        if processed_count % 10 == 0:
                            QCoreApplication.processEvents()
                for future in futures:
                    if self._is_stopped:
                        break
                    future.result()
            self.finished.emit()
        except Exception as e:
            self.error_occurred.emit(f"Critical error: {str(e)}")

# ---------------- Smart Check Thread (Non-Matching DB) ----------------
class SmartCheckThread(QThread):
    result_ready = pyqtSignal(list)
    def __init__(self, db, target_hash):
        super().__init__()
        self.db = db
        self.target_hash = target_hash
    def run(self):
        try:
            results = self.db.search_non_matching(self.target_hash)
            self.result_ready.emit(results)
        except Exception as e:
            self.result_ready.emit([])

# ---------------- Dialog: Non-Matching Hash Database ----------------
class NonMatchingDBDialog(QDialog):
    def __init__(self, db):
        super().__init__()
        self.db = db
        self.setWindowTitle("Non-Matching Hash Records")
        self.setGeometry(200, 200, 1000, 600)
        self.init_ui()
        self.apply_dialog_style()
        threading.Thread(target=self.load_data, daemon=True).start()
    def init_ui(self):
        layout = QVBoxLayout()
        self.table = QTableWidget()
        self.table.setColumnCount(10)
        self.table.setHorizontalHeaderLabels(
            ["Name", "Path", "Signature", "Status", "Size", "Type", "Created", "Modified", "Age", "Extra"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.table)
        self.setLayout(layout)
    def apply_dialog_style(self):
        self.setStyleSheet("""
            QDialog {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                              stop:0 #E0E0E0, stop:1 #F0F0F0);
                border-radius: 10px;
            }
            QTableWidget {
                background-color: #ffffff;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                selection-background-color: #D8BFD8;
                selection-color: #333;
            }
            QHeaderView::section {
                background-color: #6A5ACD;
                color: white;
                font-weight: bold;
                padding: 8px;
                border: 1px solid #6A5ACD;
                border-top-left-radius: 8px;
                border-top-right-radius: 8px;
            }
            QTableWidget::item {
                padding: 5px;
            }
            QMessageBox, QComboBox, QAbstractItemView {
                font: 10pt 'Segoe UI';
            }
        """)
    def load_data(self):
        try:
            query = '''
                SELECT file_path, file_hash FROM non_matching_hashes
                ORDER BY search_date DESC
            '''
            with self.db.conn:
                cursor = self.db.conn.execute(query)
                data = cursor.fetchall()
            self.table.setRowCount(len(data))
            for row, (path, hash_val) in enumerate(data):
                name = os.path.basename(path)
                try:
                    size = str(os.path.getsize(path))
                except Exception:
                    size = "N/A"
                try:
                    created = time.ctime(os.path.getctime(path))
                except Exception:
                    created = "N/A"
                try:
                    modified = time.ctime(os.path.getmtime(path))
                except Exception:
                    modified = "N/A"
                try:
                    age = f"{((time.time() - os.path.getctime(path)) / 86400.0):.1f} days"
                except Exception:
                    age = "N/A"
                extra = "Non-Match"
                status = "Available" if os.path.exists(path) else "Deleted"
                self.table.setItem(row, 0, QTableWidgetItem(name))
                self.table.setItem(row, 1, QTableWidgetItem(path))
                self.table.setItem(row, 2, QTableWidgetItem(hash_val))
                self.table.setItem(row, 3, QTableWidgetItem(status))
                self.table.setItem(row, 4, QTableWidgetItem(size))
                self.table.setItem(row, 5, QTableWidgetItem(os.path.splitext(path)[1]))
                self.table.setItem(row, 6, QTableWidgetItem(created))
                self.table.setItem(row, 7, QTableWidgetItem(modified))
                self.table.setItem(row, 8, QTableWidgetItem(age))
                self.table.setItem(row, 9, QTableWidgetItem(extra))
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error loading data: {str(e)}")

# ---------------- Dialog: Exclude Paths Management ----------------
class ExcludePathsDialog(QDialog):
    def __init__(self, excluded_paths):
        super().__init__()
        self.setWindowTitle("Manage Excluded Paths")
        self.setGeometry(300, 300, 500, 400)
        self.excluded_paths = excluded_paths[:]  # نسخ القائمة
        self.init_ui()
        self.apply_dialog_style()
    def init_ui(self):
        layout = QVBoxLayout()
        self.list_widget = QListWidget()
        self.list_widget.addItems(self.excluded_paths)
        layout.addWidget(self.list_widget)
        btn_layout = QHBoxLayout()
        self.btn_add = HoverButton("Add Path", icon_name="folder")
        self.btn_remove = HoverButton("Remove Selected", icon_name="clear")
        self.btn_save = HoverButton("Save", icon_name="save")
        self.btn_close = HoverButton("Close", icon_name="exit")
        btn_layout.addWidget(self.btn_add)
        btn_layout.addWidget(self.btn_remove)
        btn_layout.addWidget(self.btn_save)
        btn_layout.addStretch()
        btn_layout.addWidget(self.btn_close)
        layout.addLayout(btn_layout)
        self.setLayout(layout)
        self.btn_add.clicked.connect(self.add_path)
        self.btn_remove.clicked.connect(self.remove_selected)
        self.btn_save.clicked.connect(self.save_exclusions)
        self.btn_close.clicked.connect(self.accept)
    def apply_dialog_style(self):
        self.setStyleSheet("""
            QDialog {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                              stop:0 #E0E0E0, stop:1 #F0F0F0);
                border-radius: 10px;
            }
            QListWidget {
                background-color: #ffffff;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                padding: 5px;
            }
            QListWidget::item {
                padding: 5px;
            }
        """)
    def keyPressEvent(self, event):
        # إضافة ميزة تحديد الكل باستخدام Ctrl+A
        if event.modifiers() == Qt.ControlModifier and event.key() == Qt.Key_A:
            self.list_widget.selectAll()
        else:
            super().keyPressEvent(event)
    def add_path(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Folder to Exclude")
        if folder and folder not in self.excluded_paths:
            self.excluded_paths.append(folder)
            self.list_widget.addItem(folder)
    def remove_selected(self):
        selected_items = self.list_widget.selectedItems()
        for item in selected_items:
            self.excluded_paths.remove(item.text())
            self.list_widget.takeItem(self.list_widget.row(item))
    def save_exclusions(self):
        self.accept()
    def get_excluded_paths(self):
        return self.excluded_paths

# ---------------- Dialog: Settings (Theme, Language, Performance) ----------------
class SettingsDialog(QDialog):
    def __init__(self, current_theme="light", current_language="English"):
        super().__init__()
        self.setWindowTitle("Settings")
        self.setGeometry(400, 400, 400, 300)
        self.current_theme = current_theme
        self.current_language = current_language
        self.init_ui()
        self.apply_dialog_style()
        self.update_language_ui()
    def init_ui(self):
        layout = QVBoxLayout()
        self.lbl_theme = QLabel("Select Theme:")
        self.theme_combo = QComboBox()
        # إضافة ثيمات إضافية إلى القائمة
        self.theme_combo.addItems(["light", "dark", "Ocean Breeze", "Sunset Orange",
                                   "Midnight Purple", "Steel Gray", "Forest Green", "Ruby Red"])
        self.theme_combo.setCurrentText(self.current_theme)
        layout.addWidget(self.lbl_theme)
        layout.addWidget(self.theme_combo)
        self.lbl_language = QLabel("Select Language:")
        self.lang_combo = QComboBox()
        self.lang_combo.addItems(["English", "Arabic"])
        self.lang_combo.setCurrentText(self.current_language)
        layout.addWidget(self.lbl_language)
        layout.addWidget(self.lang_combo)
        # تغيير تسمية "Data Filter (integer)" إلى "عمر الملف"
        self.lbl_perf = QLabel("عمر الملف:")
        self.perf_input = QLineEdit()
        self.perf_input.setPlaceholderText("مثال: تاريخ الإنشاء أو آخر تعديل")
        layout.addWidget(self.lbl_perf)
        layout.addWidget(self.perf_input)
        btn_layout = QHBoxLayout()
        self.btn_ok = HoverButton("OK", icon_name="save")
        self.btn_cancel = HoverButton("Cancel", icon_name="exit")
        btn_layout.addWidget(self.btn_ok)
        btn_layout.addWidget(self.btn_cancel)
        layout.addLayout(btn_layout)
        self.setLayout(layout)
        self.btn_ok.clicked.connect(self.accept)
        self.btn_cancel.clicked.connect(self.reject)
        self.lang_combo.currentTextChanged.connect(self.update_language_ui)
    def apply_dialog_style(self):
        self.setStyleSheet("""
            QDialog {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                              stop:0 #E0E0E0, stop:1 #F0F0F0);
                border-radius: 10px;
            }
            QLabel {
                font-weight: bold;
                color: #333;
            }
            QComboBox, QLineEdit, QAbstractItemView, QMessageBox {
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 5px;
                background-color: #fff;
                font: 10pt 'Segoe UI';
            }
            HoverButton {
                background-color: #6A5ACD;
                color: white;
            }
            HoverButton:hover {
                background-color: #8A2BE2;
            }
        """)
    def update_language_ui(self):
        lang = self.lang_combo.currentText()
        if lang == "Arabic":
            self.lbl_theme.setText("اختر المظهر:")
            self.lbl_language.setText("اختر اللغة:")
            self.lbl_perf.setText("عمر الملف:")
            self.btn_ok.setText("موافق")
            self.btn_cancel.setText("إلغاء")
        else:
            self.lbl_theme.setText("Select Theme:")
            self.lbl_language.setText("Select Language:")
            self.lbl_perf.setText("File Age:")
            self.btn_ok.setText("OK")
            self.btn_cancel.setText("Cancel")
    def get_settings(self):
        return {
            "theme": self.theme_combo.currentText(),
            "language": self.lang_combo.currentText(),
            "performance": self.perf_input.text()
        }

# ---------------- Main Window with Enhanced UI, Dashboard Removed and Logs Integrated in Statistics ----------------
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.db = DatabaseManager()
        self.current_thread = None
        self.excluded_paths = []
        self.dark_mode = False
        self.disk_count = 0
        self.smart_count = 0
        # لا يتم مسح النتائج تلقائياً عند بدء بحث جديد
        self.results_data = []  # تخزين النتائج من جميع عمليات البحث
        self.log_messages = []  # سجل الأحداث
        self.init_ui()
        self.setup_connections()
        # تعيين الثيم الافتراضي بناءً على إعدادات المستخدم
        self.apply_theme("light")
        self.update_language_ui()  # تحديث النصوص بناءً على اللغة
    def init_ui(self):
        self.setWindowTitle("Professional Digital Forensic Investigation Tool")
        self.setGeometry(100, 100, 1400, 900)
        central_widget = QWidget()
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        # Top bar with glass-like effect
        top_bar = QWidget()
        top_bar.setObjectName("topBar")
        top_layout = QHBoxLayout()
        top_layout.setContentsMargins(15, 10, 15, 10)
        self.title_top = QLabel("Digital Forensic Investigation Tool")
        self.title_top.setFont(QFont("Tahoma", 18, QFont.Bold))
        self.title_top.setStyleSheet("color: white;")
        top_layout.addWidget(self.title_top)
        top_layout.addStretch()
        # زر Refresh بجانب Settings
        self.btn_refresh_top = HoverButton("Refresh All", icon_name="refresh")
        top_layout.addWidget(self.btn_refresh_top)
        self.btn_settings = HoverButton("Settings", icon_name="settings")
        top_layout.addWidget(self.btn_settings)
        self.btn_exit_top = HoverButton("Exit", icon_name="exit")
        top_layout.addWidget(self.btn_exit_top)
        top_bar.setLayout(top_layout)
        main_layout.addWidget(top_bar)
        # Splitter for Control Panel and Results/Analysis Panel (Dashboard and Logs tabs حذفت)
        splitter = QSplitter(Qt.Horizontal)
        splitter.setContentsMargins(10, 10, 10, 10)
        # Left Panel - Control Panel with glassmorphism
        self.control_panel = QWidget()
        self.control_panel.setObjectName("controlPanel")
        ctrl_layout = QVBoxLayout()
        ctrl_layout.setContentsMargins(15, 15, 15, 15)
        ctrl_layout.setSpacing(15)
        # Card 1: Search Settings
        self.search_settings_card = QGroupBox("Search Settings")
        self.search_settings_card.setObjectName("searchSettingsCard")
        ss_layout = QGridLayout()
        ss_layout.setSpacing(10)
        ss_layout.addWidget(QLabel("SHA-256:"), 0, 0)
        self.input_hash = QLineEdit()
        self.input_hash.setPlaceholderText("Enter SHA-256 hash")
        ss_layout.addWidget(self.input_hash, 0, 1)
        self.btn_calculate = HoverButton("Calculate Hash", icon_name="hash")
        ss_layout.addWidget(self.btn_calculate, 0, 2)
        ss_layout.addWidget(QLabel("Search Folder:"), 1, 0)
        self.input_folder = QLineEdit()
        self.input_folder.setPlaceholderText("Select folder to scan")
        ss_layout.addWidget(self.input_folder, 1, 1)
        self.btn_browse_folder = HoverButton("Browse", icon_name="folder")
        ss_layout.addWidget(self.btn_browse_folder, 1, 2)
        ss_layout.addWidget(QLabel("File Extensions:"), 2, 0)
        self.combo_extensions = QComboBox()
        self.combo_extensions.addItems(["all", ".txt", ".pdf", ".docx", ".jpg", ".exe", ".sys"])
        ss_layout.addWidget(self.combo_extensions, 2, 1)
        ss_layout.addWidget(QLabel("Minimum File Size (bytes):"), 3, 0)
        self.input_min_size = QLineEdit()
        self.input_min_size.setPlaceholderText("0")
        ss_layout.addWidget(self.input_min_size, 3, 1)
        # تغيير تسمية الفلتر إلى "عمر الملف"
        ss_layout.addWidget(QLabel("Age File:"), 4, 0)
        self.input_data_filter = QLineEdit()
        self.input_data_filter.setPlaceholderText("مثال: تاريخ الإنشاء أو آخر تعديل")
        self.input_data_filter.setValidator(QIntValidator())
        ss_layout.addWidget(self.input_data_filter, 4, 1)
        ss_layout.addWidget(QLabel("Digital Signature:"), 5, 0)
        self.combo_signature = QComboBox()
        self.combo_signature.addItems(["Valid", "Invalid", "Unknown", "All"])
        ss_layout.addWidget(self.combo_signature, 5, 1)
        self.search_settings_card.setLayout(ss_layout)
        ctrl_layout.addWidget(self.search_settings_card)
        # Card 2: Quick Actions (زر Resume محذوف)
        self.quick_actions_card = QGroupBox("Quick Actions")
        self.quick_actions_card.setObjectName("quickActionsCard")
        qa_layout = QHBoxLayout()
        self.btn_normal_search = HoverButton("Normal Search", icon_name="search")
        self.btn_smart_search = HoverButton("Smart Search", icon_name="search")
        # زر Pause متبقي، زر Resume محذوف
        self.btn_pause = HoverButton("Stop", icon_name="stop")
        qa_layout.addWidget(self.btn_normal_search)
        qa_layout.addWidget(self.btn_smart_search)
        qa_layout.addWidget(self.btn_pause)
        self.quick_actions_card.setLayout(qa_layout)
        ctrl_layout.addWidget(self.quick_actions_card)
        # Progress bar with dynamic status text
        self.progress_label = QLabel("Search Progress: Not started")
        self.progress_label.setFont(QFont("Arial", 10, QFont.Bold))
        self.progress_bar = AnimatedProgressBar()
        ctrl_layout.addWidget(self.progress_label)
        ctrl_layout.addWidget(self.progress_bar)
        # Exclude paths and Non-Matching DB buttons
        btns_layout = QHBoxLayout()
        self.btn_exclude_paths = HoverButton("Manage Excluded Paths", icon_name="exclude")
        self.btn_non_matching_db = HoverButton("Non-Matching Database", icon_name="database")
        btns_layout.addWidget(self.btn_exclude_paths)
        btns_layout.addWidget(self.btn_non_matching_db)
        ctrl_layout.addLayout(btns_layout)
        # Card 3: Statistics with integrated Log area (الملاحظات الصغيرة)
        self.statistics_card = QGroupBox("Statistics")
        self.statistics_card.setObjectName("statisticsCard")
        stats_layout = QVBoxLayout()
        self.chart_widget = ContemporaryChartWidget()
        stats_layout.addWidget(self.chart_widget)
        self.label_total_files = QLabel("Total Files: 0")
        self.label_matches = QLabel("Matches: 0")
        self.label_speed = QLabel("Scan Speed: N/A")
        stats_layout.addWidget(self.label_total_files)
        stats_layout.addWidget(self.label_matches)
        stats_layout.addWidget(self.label_speed)
        self.statistics_card.setLayout(stats_layout)
        ctrl_layout.addWidget(self.statistics_card)
        ctrl_layout.addStretch()
        self.control_panel.setLayout(ctrl_layout)
        splitter.addWidget(self.control_panel)
        # Right Panel - Tab Widget for Results and Analysis (Dashboard and Logs removed)
        self.results_panel = QTabWidget()
        self.results_panel.setObjectName("resultsPanel")
        # Tab 1: Search Results (نتائج البحث تظل كما هي)
        self.tab_results = QWidget()
        results_layout = QVBoxLayout()
        self.results_table = QTableWidget()
        self.results_table.setColumnCount(10)
        self.results_table.setHorizontalHeaderLabels(
            ["Name", "Path", "Signature", "Status", "Size", "Type", "Created", "Modified", "Age", "Extra"])
        self.results_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.results_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.results_table.setAlternatingRowColors(True)
        results_layout.addWidget(self.results_table)
        self.btn_clear_results = HoverButton("Clear Results", icon_name="clear")
        refresh_layout = QHBoxLayout()
        refresh_layout.addWidget(self.btn_clear_results)
        refresh_layout.addStretch()
        results_layout.addLayout(refresh_layout)
        self.tab_results.setLayout(results_layout)
        self.results_panel.addTab(self.tab_results, "Search Results")
        # Tab 2: Analysis
        self.tab_analysis = QWidget()
        analysis_layout = QVBoxLayout()
        self.analysis_text = QTextEdit()
        self.analysis_text.setReadOnly(True)
        analysis_layout.addWidget(self.analysis_text)
        analysis_btn_layout = QHBoxLayout()
        self.btn_generate_report = HoverButton("Generate Report", icon_name="report")
        self.btn_save_report = HoverButton("Save Report", icon_name="save")
        analysis_btn_layout.addWidget(self.btn_generate_report)
        analysis_btn_layout.addWidget(self.btn_save_report)
        analysis_layout.addLayout(analysis_btn_layout)
        self.tab_analysis.setLayout(analysis_layout)
        self.results_panel.addTab(self.tab_analysis, "Analysis")
        splitter.addWidget(self.results_panel)
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 2)
        main_layout.addWidget(splitter)
        # Status Bar with dynamic status indicator
        self.status_bar = self.statusBar()
        self.status_indicator = QLabel("●")
        self.status_indicator.setStyleSheet("color: green; font-size:16px;")
        self.status_text = QLabel("Ready")
        self.status_progress = QProgressBar()
        self.status_progress.setMaximumWidth(150)
        self.status_progress.setRange(0, 100)
        self.status_bar.addPermanentWidget(self.status_indicator)
        self.status_bar.addPermanentWidget(self.status_text)
        self.status_bar.addPermanentWidget(self.status_progress)
        self.setCentralWidget(central_widget)
        self.animate_group_boxes()
        # Connect double-click on results to show details
        self.results_table.cellDoubleClicked.connect(self.show_result_details)
        # Install scroll event filter for infinite scrolling simulation on results_table
        self.results_table.verticalScrollBar().valueChanged.connect(self.check_infinite_scroll)
    def animate_group_boxes(self):
        for group_box in [self.search_settings_card, self.quick_actions_card, self.statistics_card]:
            group_box.setWindowOpacity(0)
            animation = QPropertyAnimation(group_box, b"windowOpacity")
            animation.setDuration(500)
            animation.setStartValue(0.0)
            animation.setEndValue(1.0)
            animation.setEasingCurve(QEasingCurve.InQuad)
            animation.start()
    def setup_connections(self):
        self.btn_browse_folder.clicked.connect(self.browse_folder)
        self.btn_calculate.clicked.connect(self.calculate_hash)
        self.btn_normal_search.clicked.connect(self.start_normal_search)
        self.btn_smart_search.clicked.connect(self.start_smart_search)
        self.btn_pause.clicked.connect(self.stop_search)
        self.btn_clear_results.clicked.connect(self.clear_results)
        self.btn_exclude_paths.clicked.connect(self.manage_excluded_paths)
        self.btn_non_matching_db.clicked.connect(self.open_non_matching_db)
        self.btn_settings.clicked.connect(self.open_settings)
        self.btn_exit_top.clicked.connect(self.exit_program)
        self.btn_refresh_top.clicked.connect(self.refresh_all)
        self.btn_save_report.clicked.connect(self.save_report_and_show_excluded)
        self.btn_generate_report.clicked.connect(self.generate_report)
    def log_event(self, message):
        timestamp = datetime.datetime.now().strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}"
        self.log_messages.append(log_entry)
        # أيضًا نقوم بتحديث منطقة السجل في لوحة الإحصائيات
        self.chart_widget.update_log(log_entry)
    def apply_theme(self, theme):
        self.dark_mode = theme.lower() in ["dark", "midnight purple", "steel gray", "forest green", "ruby red"]
        if theme.lower() == "ocean breeze":
            qss = """
                QMainWindow { background: #DEF; }
                #topBar { background-color: #4FC3F7; }
                QLabel { color: #333; }
                QGroupBox { background: rgba(255,255,255,0.7); border: 1px solid #4FC3F7; border-radius: 12px; }
                QLineEdit, QComboBox, QAbstractItemView, QMessageBox { background: rgba(227,242,253,0.8); color: #333; }
            """
            prog_colors = ("#4FC3F7", "#0288D1")
        elif theme.lower() == "sunset orange":
            qss = """
                QMainWindow { background: #FFE0B2; }
                #topBar { background-color: #FF9800; }
                QLabel { color: #333; }
                QGroupBox { background: rgba(255,255,255,0.7); border: 1px solid #FF9800; border-radius: 12px; }
                QLineEdit, QComboBox, QAbstractItemView, QMessageBox { background: rgba(255,243,224,0.8); color: #333; }
            """
            prog_colors = ("#FF9800", "#F57C00")
        elif theme.lower() == "midnight purple":
            qss = """
                QMainWindow { background: #2E003E; }
                #topBar { background-color: #4A148C; }
                QLabel { color: #EEE; }
                QGroupBox { background: rgba(62,31,91,0.8); border: 1px solid #4A148C; border-radius: 12px; }
                QLineEdit, QComboBox, QAbstractItemView, QMessageBox { background: #4A148C; color: #EEE; }
            """
            prog_colors = ("#4A148C", "#6A1B9A")
        elif theme.lower() == "steel gray":
            qss = """
                QMainWindow { background: #ECEFF1; }
                #topBar { background-color: #607D8B; }
                QLabel { color: #333; }
                QGroupBox { background: rgba(255,255,255,0.7); border: 1px solid #607D8B; border-radius: 12px; }
                QLineEdit, QComboBox, QAbstractItemView, QMessageBox { background: #CFD8DC; color: #333; }
            """
            prog_colors = ("#607D8B", "#455A64")
        elif theme.lower() == "forest green":
            qss = """
                QMainWindow { background: #E8F5E9; }
                #topBar { background-color: #388E3C; }
                QLabel { color: #333; }
                QGroupBox { background: rgba(255,255,255,0.7); border: 1px solid #388E3C; border-radius: 12px; }
                QLineEdit, QComboBox, QAbstractItemView, QMessageBox { background: #C8E6C9; color: #333; }
            """
            prog_colors = ("#388E3C", "#2E7D32")
        elif theme.lower() == "ruby red":
            qss = """
                QMainWindow { background: #FCE4EC; }
                #topBar { background-color: #D81B60; }
                QLabel { color: #333; }
                QGroupBox { background: rgba(255,255,255,0.7); border: 1px solid #D81B60; border-radius: 12px; }
                QLineEdit, QComboBox, QAbstractItemView, QMessageBox { background: #F8BBD0; color: #333; }
            """
            prog_colors = ("#D81B60", "#AD1457")
        elif theme.lower() in ["dark", "ruby red", "midnight purple", "steel gray", "forest green"]:
            qss = """
                QMainWindow { background: #2c2c2c; }
                #topBar { background-color: #1a1a1a; border-bottom: 1px solid #333; }
                QLabel { color: #e0e0e0; }
                QGroupBox { background: rgba(60,60,60,0.8); border: 1px solid #444; border-radius: 12px; padding: 10px; color: #e0e0e0; }
                QLineEdit, QComboBox, QAbstractItemView, QMessageBox { background: #555; color: #e0e0e0; padding: 8px; border: 1px solid #666; border-radius: 8px; }
            """
            prog_colors = ("#6A5ACD", "#9370DB")
        else:
            qss = """
                QMainWindow { background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #F0F0F0, stop:0.5 #E0E0E0, stop:1 #D0D0D0); }
                #topBar { background-color: #6A5ACD; border-bottom: 1px solid #5B3C9B; }
                QLabel { color: #333; }
                QGroupBox { background: rgba(255,255,255,0.7); border: 1px solid #ccc; border-radius: 12px; padding: 10px; }
                QLineEdit, QComboBox, QAbstractItemView, QMessageBox { background: #ffffff; color: #333; padding: 8px; border: 1px solid #ccc; border-radius: 8px; }
            """
            prog_colors = ("#6A5ACD", "#8A2BE2")
        self.setStyleSheet(qss)
        self.progress_bar.set_gradient_colors(prog_colors[0], prog_colors[1])
    def update_language_ui(self):
        current_lang = self.get_current_language()
        if current_lang == "Arabic":
            self.title_top.setText("أداة التحقيق الجنائي الرقمي")
        else:
            self.title_top.setText("ForensicX")
    def browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Search Folder")
        if folder:
            self.input_folder.setText(folder)
    def calculate_hash(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Select a File")
        if file_path:
            try:
                hasher = hashlib.sha256()
                with open(file_path, 'rb') as f:
                    while True:
                        chunk = f.read(131072)
                        if not chunk:
                            break
                        hasher.update(chunk)
                hash_val = hasher.hexdigest()
                self.input_hash.setText(hash_val)
                QMessageBox.information(self, "Success", f"File Hash:\n{hash_val}")
                self.log_event("Calculated hash for file: " + file_path)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Unable to read file: {str(e)}")
    def start_normal_search(self):
        self.status_text.setText("Working")
        self.status_indicator.setStyleSheet("color: orange; font-size:16px;")
        target_hash = self.input_hash.text().strip().lower()
        folder = self.input_folder.text()
        if len(target_hash) != 64:
            QMessageBox.warning(self, "Error", "Hash must be 64 characters long")
            return
        if not os.path.isdir(folder):
            QMessageBox.warning(self, "Error", "Invalid search folder")
            return
        try:
            min_size = int(self.input_min_size.text()) if self.input_min_size.text() else 0
        except ValueError:
            QMessageBox.warning(self, "Error", "Minimum file size must be a number")
            return
        data_filter = int(self.input_data_filter.text()) if self.input_data_filter.text() else None
        digital_signature = self.combo_signature.currentText()
        extensions = [self.combo_extensions.currentText()]
        self.progress_bar.setRange(0, 0)
        self.status_progress.setRange(0, 0)
        # لا يتم مسح النتائج القديمة، لذا لا نقوم بتهيئة self.results_data أو استدعاء clear_results()
        self.log_event("Starting normal search")
        self.current_thread = LocalSearchThread([folder], target_hash, extensions, self.excluded_paths, min_size,
                                                data_filter, digital_signature)
        self.current_thread.result_found.connect(self.handle_result_found)
        self.current_thread.error_occurred.connect(lambda e: QMessageBox.critical(self, "Error", e))
        self.current_thread.finished.connect(self.search_finished)
        self.current_thread.start()
        self.progress_label.setText("Search Progress: Scanning...")
    def start_smart_search(self):
        self.status_text.setText("Working (Smart)")
        self.status_indicator.setStyleSheet("color: orange; font-size:16px;")
        target_hash = self.input_hash.text().strip().lower()
        folder = self.input_folder.text()
        if len(target_hash) != 64:
            QMessageBox.warning(self, "Error", "Hash must be 64 characters long")
            return
        if not os.path.isdir(folder):
            QMessageBox.warning(self, "Error", "Invalid search folder")
            return
        self.log_event("Starting smart search")
        self.current_thread = None
        self.smart_thread = SmartCheckThread(self.db, target_hash)
        self.smart_thread.result_ready.connect(lambda results: self.handle_smart_results(results, target_hash, folder))
        self.smart_thread.start()
        self.progress_label.setText("Search Progress: Smart scanning...")
    def handle_smart_results(self, results, target_hash, folder):
        if results:
            self.smart_count = len(results)
            for path, hash_val in results:
                row_data = {
                    "name": os.path.basename(path),
                    "path": path,
                    "signature": hash_val,
                    "status": "Available" if os.path.exists(path) else "Deleted",
                    "size": str(os.path.getsize(path)) if os.path.exists(path) else "N/A",
                    "type": os.path.splitext(path)[1],
                    "created": time.ctime(os.path.getctime(path)) if os.path.exists(path) else "N/A",
                    "modified": time.ctime(os.path.getmtime(path)) if os.path.exists(path) else "N/A",
                    "age": f"{((time.time() - os.path.getctime(path)) / 86400.0):.1f} days" if os.path.exists(path) else "N/A",
                    "extra": "Smart"
                }
                self.results_data.append(row_data)
                self.add_result_row(path, hash_val, "Smart", is_match=True)
            self.chart_widget.update_chart(self.disk_count, self.smart_count)
            self.progress_label.setText("Search Progress: Smart search successful.")
            self.status_text.setText("Success")
            self.status_indicator.setStyleSheet("color: green; font-size:16px;")
            self.log_event("Smart search completed successfully")
        else:
            reply = QMessageBox.question(self, "No Result",
                                         "No matching record found in DB. Do you want to start a normal search?",
                                         QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                self.start_normal_search()
            else:
                self.progress_label.setText("Search Progress: Idle")
                self.status_text.setText("Ready")
                self.status_indicator.setStyleSheet("color: green; font-size:16px;")
                self.log_event("Smart search found no results; user opted not to run normal search")
    def get_file_icon(self, file_path):
        return get_icon("file", self)
    def handle_result_found(self, path, hash_val):
        self.disk_count += 1
        row_data = {
            "name": os.path.basename(path),
            "path": path,
            "signature": hash_val,
            "status": "Available" if os.path.exists(path) else "Deleted",
            "size": str(os.path.getsize(path)) if os.path.exists(path) else "N/A",
            "type": os.path.splitext(path)[1],
            "created": time.ctime(os.path.getctime(path)) if os.path.exists(path) else "N/A",
            "modified": time.ctime(os.path.getmtime(path)) if os.path.exists(path) else "N/A",
            "age": f"{((time.time() - os.path.getctime(path)) / 86400.0):.1f} days" if os.path.exists(path) else "N/A",
            "extra": "Normal"
        }
        self.results_data.append(row_data)
        self.add_result_row(path, hash_val, "Normal", is_match=True)
        total_files = self.disk_count + self.smart_count
        self.label_total_files.setText(f"Total Files: {total_files}")
        self.label_matches.setText(f"Matches: {self.disk_count}")
        self.label_speed.setText("Scan Speed: Calculating...")
        self.chart_widget.update_chart(self.disk_count, self.smart_count)
    def add_result_row(self, path, hash_val, source, is_match=False):
        row_pos = self.results_table.rowCount()
        self.results_table.insertRow(row_pos)
        name = os.path.basename(path)
        try:
            size = str(os.path.getsize(path))
        except Exception:
            size = "N/A"
        try:
            created = time.ctime(os.path.getctime(path))
        except Exception:
            created = "N/A"
        try:
            modified = time.ctime(os.path.getmtime(path))
        except Exception:
            modified = "N/A"
        try:
            age = f"{((time.time() - os.path.getctime(path)) / 86400.0):.1f} days"
        except Exception:
            age = "N/A"
        extra = source
        status = "Available" if os.path.exists(path) else "Deleted"
        items = [name, path, hash_val, status, size,
                 os.path.splitext(path)[1], created, modified, age, extra]
        for col, val in enumerate(items):
            item = QTableWidgetItem(val)
            if col == 0:
                item.setIcon(self.get_file_icon(path))
            if is_match:
                item.setBackground(QColor("#D8BFD8"))
                item.setForeground(QColor("#333333"))
            self.results_table.setItem(row_pos, col, item)
        # إظهار تنبيه عند وصول عدد المطابقات إلى مضاعفات معينة
        if is_match and source == "Normal" and self.disk_count % 5 == 0:
            QMessageBox.information(self, "Important", "A significant number of matches have been found!")
    def show_result_details(self, row, column):
        if row < len(self.results_data):
            data = self.results_data[row]
            details = f"<b>Name:</b> {data.get('name', '')}<br>"
            details += f"<b>Path:</b> {data.get('path', '')}<br>"
            details += f"<b>Signature:</b> {data.get('signature', '')}<br>"
            details += f"<b>Status:</b> {data.get('status', '')}<br>"
            details += f"<b>Size:</b> {data.get('size', '')}<br>"
            details += f"<b>Type:</b> {data.get('type', '')}<br>"
            details += f"<b>Created:</b> {data.get('created', '')}<br>"
            details += f"<b>Modified:</b> {data.get('modified', '')}<br>"
            details += f"<b>Age:</b> {data.get('age', '')}<br>"
            details += f"<b>Extra:</b> {data.get('extra', '')}<br>"
            QMessageBox.information(self, "Result Details", details)
    def search_finished(self):
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(100)
        self.status_progress.setRange(0, 100)
        self.status_progress.setValue(100)
        self.progress_label.setText("Search Progress: Completed")
        self.status_text.setText("Success")
        self.status_indicator.setStyleSheet("color: green; font-size:16px;")
        self.log_event("Search finished successfully")
    def stop_search(self):
        if self.current_thread:
            self.current_thread.stop()
            self.progress_label.setText("Search Progress: Stopped")
            self.status_text.setText("Stopped")
            self.status_indicator.setStyleSheet("color: blue; font-size:16px;")
            self.log_event("Search stopped by user")
    def clear_results(self):
        # يسمح للمستخدم بمسح النتائج يدوياً، لكن البحث الجديد لا يمسح النتائج القديمة تلقائياً
        self.results_data = []
        self.results_table.setRowCount(0)
        self.chart_widget.update_chart(0, 0)
        self.label_total_files.setText("Total Files: 0")
        self.label_matches.setText("Matches: 0")
        self.label_speed.setText("Scan Speed: N/A")
        self.progress_label.setText("Search Progress: Cleared")
        self.disk_count = 0
        self.smart_count = 0
        self.log_event("Results cleared")
    def refresh_all(self):
        self.clear_results()
        QMessageBox.information(self, "Refresh", "All results cleared. Please start a new search.")
        self.log_event("Refreshed all results")
    def manage_excluded_paths(self):
        dialog = ExcludePathsDialog(self.excluded_paths)
        if dialog.exec_():
            self.excluded_paths = dialog.get_excluded_paths()
            self.log_event("Updated excluded paths")
    def open_non_matching_db(self):
        try:
            dialog = NonMatchingDBDialog(self.db)
            dialog.exec_()
            self.log_event("Opened non-matching database")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open Non-Matching Database:\n{str(e)}")
    def open_settings(self):
        dialog = SettingsDialog("dark" if self.dark_mode else "light", self.get_current_language())
        if dialog.exec_():
            settings = dialog.get_settings()
            self.apply_theme(settings["theme"])
            self.update_language_ui()
            self.log_event("Settings updated")
    def get_current_language(self):
        return "Arabic" if self.title_top.text() == "أداة التحقيق الجنائي الرقمي" else "English"
    def generate_report(self):
        report = "<h2>Detailed Analysis Report</h2>"
        report += f"<p><b>Total Files Scanned:</b> {self.label_total_files.text().split(':')[-1].strip()}</p>"
        report += f"<p><b>Matches Found:</b> {self.label_matches.text().split(':')[-1].strip()}</p>"
        report += f"<p><b>Scan Speed:</b> {self.label_speed.text().split(':')[-1].strip()}</p>"
        report += "<hr>"
        for row in self.results_data:
            report += f"<p><b>{row.get('name')}</b><br>"
            report += f"Path: {row.get('path')}<br>"
            report += f"Signature: {row.get('signature')}<br>"
            report += f"Status: {row.get('status')}<br>"
            report += f"Size: {row.get('size')}<br>"
            report += f"Type: {row.get('type')}<br>"
            report += f"Created: {row.get('created')}<br>"
            report += f"Modified: {row.get('modified')}<br>"
            report += f"Age: {row.get('age')}<br>"
            report += f"Extra: {row.get('extra')}</p><hr>"
        self.analysis_text.setHtml(report)
        self.log_event("Generated analysis report")
    def save_report(self):
        investigator, ok = QInputDialog.getText(self, "Investigator", "Enter Investigator Name:")
        if not ok:
            return
        date_str = QInputDialog.getText(self, "Investigation Date", "Enter Investigation Date (YYYY-MM-DD):")[0]
        formats = ["PDF", "JSON", "CSV", "XLSX", "Word"]
        fmt, ok = QInputDialog.getItem(self, "Select Format", "Select save format:", formats, 0, False)
        if not ok or not fmt:
            return
        file_filter = ""
        if fmt == "PDF":
            file_filter = "PDF Files (*.pdf)"
        elif fmt == "JSON":
            file_filter = "JSON Files (*.json)"
        elif fmt == "CSV":
            file_filter = "CSV Files (*.csv)"
        elif fmt == "XLSX":
            file_filter = "Excel Files (*.xlsx)"
        elif fmt == "Word":
            file_filter = "Word Documents (*.docx)"
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Report", "", file_filter)
        if not file_path:
            return
        report_content = self.analysis_text.toHtml()
        try:
            if fmt == "PDF":
                c = canvas.Canvas(file_path, pagesize=letter)
                textobject = c.beginText()
                textobject.setTextOrigin(letter[0] * 0.1, letter[1] * 0.9)
                textobject.setFont("Helvetica", 10)
                plain_text = self.analysis_text.toPlainText()
                for line in plain_text.split('\n'):
                    textobject.textLine(line)
                c.drawText(textobject)
                c.save()
            elif fmt == "JSON":
                with open(file_path, "w", encoding="utf-8") as f:
                    json.dump(self.results_data, f, ensure_ascii=False, indent=4)
            elif fmt == "CSV":
                with open(file_path, "w", newline='', encoding="utf-8") as f:
                    writer = csv.DictWriter(f, fieldnames=self.results_data[0].keys())
                    writer.writeheader()
                    writer.writerows(self.results_data)
            elif fmt == "XLSX":
                if openpyxl is None:
                    QMessageBox.warning(self, "Dependency Missing", "openpyxl is required for XLSX export. Please install it via pip.")
                    return
                wb = openpyxl.Workbook()
                ws = wb.active
                headers = list(self.results_data[0].keys())
                ws.append(headers)
                for row in self.results_data:
                    ws.append(list(row.values()))
                wb.save(file_path)
            elif fmt == "Word":
                document = Document()
                document.add_heading("Detailed Analysis Report", level=2)
                document.add_paragraph(f"Investigator: {investigator}")
                document.add_paragraph(f"Investigation Date: {date_str}")
                document.add_paragraph(f"Total Files Scanned: {self.label_total_files.text().split(':')[-1].strip()}")
                document.add_paragraph(f"Matches Found: {self.label_matches.text().split(':')[-1].strip()}")
                document.add_paragraph(f"Scan Speed: {self.label_speed.text().split(':')[-1].strip()}")
                for row in self.results_data:
                    document.add_heading(row.get("name"), level=3)
                    p = document.add_paragraph()
                    p.add_run(f"Path: {row.get('path')}\n")
                    p.add_run(f"Signature: {row.get('signature')}\n")
                    p.add_run(f"Status: {row.get('status')}\n")
                    p.add_run(f"Size: {row.get('size')}\n")
                    p.add_run(f"Type: {row.get('type')}\n")
                    p.add_run(f"Created: {row.get('created')}\n")
                    p.add_run(f"Modified: {row.get('modified')}\n")
                    p.add_run(f"Age: {row.get('age')}\n")
                    p.add_run(f"Extra: {row.get('extra')}\n")
                    document.add_page_break()
                document.save(file_path)
            QMessageBox.information(self, "Success", "Report saved successfully.")
            self.log_event("Report saved to: " + file_path)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save report: {str(e)}")
            self.log_event("Error saving report: " + str(e))
    def save_report_and_show_excluded(self):
        self.save_report()
        self.manage_excluded_paths()
    def exit_program(self):
        if self.current_thread and self.current_thread.isRunning():
            self.current_thread.stop()
            self.current_thread.wait()
        self.log_event("Exiting application")
        QApplication.quit()
    def refresh_results(self):
        self.results_table.setRowCount(0)
        for entry in self.results_data:
            self.add_result_row(entry.get("path"), entry.get("signature"), entry.get("extra"), is_match=True)
        QMessageBox.information(self, "Refresh", "Results refreshed!")
        self.log_event("Refreshed results table")
    def check_infinite_scroll(self, value):
        # When scrollbar reaches maximum, simulate lazy load (no extra data available in this demo)
        scroll_bar = self.results_table.verticalScrollBar()
        if value == scroll_bar.maximum():
            self.log_event("Reached bottom – infinite scrolling simulation (no more data to load)")
    def closeEvent(self, event):
        if self.current_thread and self.current_thread.isRunning():
            self.current_thread.stop()
            self.current_thread.wait()
        event.accept()

# ---------------- Main Execution ----------------
if __name__ == "__main__":
    app = QApplication(sys.argv)
    splash = SplashScreen()
    splash.show()
    app.processEvents()
    splash.exec_()
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())